# -*- coding: utf-8 -*-
"""Central version registry for the human-facing documents.

ONE place to version every generated doc, so bumping stays single-source
(no hand-editing each cover). When you regenerate a doc with CHANGED content,
bump its version + date here; the generators stamp it automatically onto the
cover and into a per-page footer (with page numbers).

Scheme: simple MAJOR.MINOR per doc.
  - MINOR bump  = content correction / added matchup / re-tuned sim.
  - MAJOR bump  = a new event, a new list, or a structural rewrite of the doc.
Dates are the CONTENT date (only changes when you bump) — reproducible, not
wall-clock, so a clean rebuild produces an identical file.

The version is also embedded in the OUTPUT FILENAME (e.g. LSO-Runbook-v1.0.docx),
so bumping a version writes a NEW file and leaves the old revision on disk. Use
all_stale() to find superseded revisions on disk when you want to prune.
"""
import os

# key -> (version, YYYY-MM-DD content date, human title, one-line note, versionless output path)
DOCS = {
    # --- RETIRED 2026-07-30: List-A / pre-field Knights LSO docs, superseded by the
    #     death_rnr LSO FIELD set (lso-field-*). Kept in _superseded/ for reference. ---
    "lso-runbook":    ("1.3", "2026-07-27", "LSO Runbook — Imperial Knights",
                       "Meta add: Imperial Knights MIRROR (2Cas/1Lancer + 1Cas/2Lancer); SW Beastslayer; OC34.",
                       "docs/reports/_superseded/knights-list-a/LSO-Runbook.pdf"),
    "lso-analysis":   ("1.3", "2026-07-27", "LSO Knights — List & Analysis",
                       "Meta add: Knights mirror matchups; SW Beastslayer; List A locked at 1,970.",
                       "docs/reports/_superseded/knights-list-a/LSO-Knights-List-and-Analysis.xlsx"),
    "lso-decision":   ("1.3", "2026-07-27", "LSO Knights — List Decision",
                       "Meta add: Knights mirror (best=B); SW Beastslayer; List A vs B, prevalence-weighted.",
                       "docs/reports/_superseded/knights-list-a/LSO-Knights-List-Decision.xlsx"),
    "gv-runbook":     ("2.0", "2026-07-27", "Great Value — LSO Runbook",
                       "Sim now models GV board control from explicit OC bricks (OC34 + losable OC21 cyclone).",
                       "docs/reports/great-value/GV-LSO-Runbook.pdf"),
    "gv-analysis":    ("2.0", "2026-07-27", "Great Value — LSO Analysis",
                       "Sim now models GV board control from explicit OC bricks (OC34 + losable OC21 cyclone).",
                       "docs/reports/great-value/GV-LSO-Analysis.xlsx"),
    "gv-sim":         ("1.1", "2026-07-27", "Great Value vs Knights — Full-Game Simulation",
                       "OC34 brick fix (out-hold needs OC35+); hardened sim ~60/40 Knights.",
                       "docs/reports/great-value/Great-Value-vs-Knights-Full-Game-Simulation.pdf"),
    "sisters-battle": ("1.0", "2026-07-27", "Adepta Sororitas — Battle Plan",
                       "Disruption lock; all points MFM, rules 11E.",
                       "docs/reports/sisters/Sisters-Battle-Plan.docx"),
    "sisters-qref":   ("1.0", "2026-07-27", "Adepta Sororitas — Tabletop Quick Reference",
                       "Per-matchup mission/quick-ref.",
                       "docs/reports/sisters/Sisters-Quick-Reference.docx"),
    "bt-runbook":     ("2.0", "2026-07-27", "Black Templars — 'Send Help' FIXED Runbook",
                       "List promoted to v2b (2x Land Raider Crusader, Assault Ramp delivery) + version history.",
                       "docs/reports/black-templars/BT-SendHelp-Fixed-Runbook.pdf"),
    "bt-analysis":    ("2.0", "2026-07-27", "Black Templars — 'Send Help' FIXED Analysis",
                       "v2b promoted; 10 matchups + sim + v1/v2a/v2b version history.",
                       "docs/reports/black-templars/BT-SendHelp-Fixed-Analysis.xlsx"),
    "bt-bastion-runbook":  ("2.0", "2026-07-27", "Black Templars — Templar Bastion Runbook",
                            "The friend's Bastion Task Force list + tapestry + per-archetype battle plans.",
                            "docs/reports/black-templars/BT-Templar-Bastion-Runbook.pdf"),
    "bt-bastion-analysis": ("2.0", "2026-07-27", "Black Templars — Templar Bastion Analysis",
                            "10 matchups + data-driven sim (~55%, durable hold-and-grind).",
                            "docs/reports/black-templars/BT-Templar-Bastion-Analysis.xlsx"),
    "custodes-analysis": ("1.0", "2026-07-28", "Adeptus Custodes — 'The Better Thing 2' Analysis",
                          "Meta Slayers list; real-mission sim (Priority Assets disposition); Purge-the-Foe is the bogey, ~52%.",
                          "docs/reports/custodes/Custodes-Better-Thing-2-Analysis.xlsx"),
    "custodes-runbook": ("1.0", "2026-07-28", "Adeptus Custodes — 'The Better Thing 2' Runbook",
                         "Full tapestry + per-matchup battle plans keyed to the forced mission; PDF-only pipeline.",
                         "docs/reports/custodes/Custodes-Better-Thing-2-Runbook.pdf"),
    "knights-analysis": ("2.0", "2026-07-28", "Imperial Knights — LSO List A Analysis",
                         "Rebuilt on the real-mission engine (Purge the Foe, 10k games): great vs Priority-Assets, poor vs Take-and-Hold; ~46%.",
                         "docs/reports/_superseded/knights-list-a/Knights-List-A-Analysis.xlsx"),
    "knights-runbook": ("2.0", "2026-07-28", "Imperial Knights — LSO List A Runbook",
                        "Per-matchup plans keyed to the FORCED mission; the disposition-driven scoring split is the headline.",
                        "docs/reports/_superseded/knights-list-a/Knights-List-A-Runbook.pdf"),
    # --- LSO 2026 FIELD deliverables: death_rnr vs the real 324-list BCP field (by archetype) ---
    "lso-field-dossier":  ("2.0", "2026-07-30", "LSO 2026 — Field Dossier (Death by Rock and Roll)",
                           "death_rnr vs the REAL 324-list LSO field, by archetype; each runbook has PRIORITY "
                           "KILLS (damage) + BOARD LYNCHPINS (OC/board-control).",
                           "docs/reports/knights/death_rnr-LSO-Field-Dossier.pdf"),
    "lso-field-analysis": ("2.0", "2026-07-30", "LSO 2026 — Field Analysis",
                           "Matchup Map + Threats & Plan over the real field; verdict-coloured, Top-OC holder + "
                           "board lynchpins per archetype.",
                           "docs/reports/knights/death_rnr-LSO-Field-Analysis.xlsx"),
    "lso-field-guide":    ("2.0", "2026-07-30", "LSO 2026 — Interactive Field Guide",
                           "Enemy-pilot lookup + verdict-filtered threat map + runbooks; self-contained HTML "
                           "(also the published Artifact).",
                           "docs/reports/knights/death_rnr-LSO-Field-Guide.html"),
    "lso-field-roster":   ("2.0", "2026-07-30", "LSO 2026 — Roster (linked)",
                           "All 330 players, each name linked to their BCP list; filterable by faction.",
                           "docs/reports/knights/death_rnr-LSO-Roster.html"),
}


def ver(key):    return DOCS[key][0]
def date(key):   return DOCS[key][1]
def title(key):  return DOCS[key][2]
def note(key):   return DOCS[key][3]
def relpath(key):return DOCS[key][4]


def out_path(key):
    """Versioned output path, e.g. 'docs/.../LSO-Runbook-v1.0.docx'."""
    stem, ext = os.path.splitext(relpath(key))
    return f"{stem}-v{ver(key)}{ext}"


def all_stale():
    """Return existing on-disk files that match a doc's stem but NOT its current version
    (superseded revisions, docx or pdf), so a prune step can remove them."""
    import glob
    keep, stale = set(), []
    for k in DOCS:
        stem, ext = os.path.splitext(relpath(k))
        keep.add(out_path(k))
        keep.add(f"{stem}-v{ver(k)}.pdf")
        for f in glob.glob(f"{glob.escape(stem)}-v*{ext}") + glob.glob(f"{glob.escape(stem)}-v*.pdf"):
            if f not in keep:
                stale.append(f)
    return sorted(set(stale) - keep)


def cover_line(key):
    """Prominent version line for a doc cover, e.g. 'Version 1.0 · 2026-07-27'."""
    return f"Version {ver(key)} · {date(key)}"


def footer_text(key):
    """Compact footer string (no page number), e.g. the title + version + date."""
    return f"{title(key)}  ·  v{ver(key)}  ·  {date(key)}"


# ---- python-docx: per-page footer with 'Page N of M' -------------------------
def stamp_docx(doc, key, size=8, grey=0x80):
    """Set a centred footer on every section: '<title> · v<ver> · <date> · Page N of M'."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _field(run, instr):
        for typ, txt in (("begin", None), (None, instr), ("end", None)):
            if typ:
                fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), typ); run._r.append(fc)
            else:
                it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
                it.text = txt; run._r.append(it)

    col = RGBColor(grey, grey, grey)
    for sec in doc.sections:
        sec.footer.is_linked_to_previous = False
        p = sec.footer.paragraphs[0]
        p.text = ""; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(footer_text(key) + "  ·  Page "); r.font.size = Pt(size); r.font.color.rgb = col
        rp = p.add_run(); rp.font.size = Pt(size); rp.font.color.rgb = col; _field(rp, "PAGE")
        r2 = p.add_run(" of "); r2.font.size = Pt(size); r2.font.color.rgb = col
        rn = p.add_run(); rn.font.size = Pt(size); rn.font.color.rgb = col; _field(rn, "NUMPAGES")


# ---- openpyxl: print header/footer (both render in the PDF export) ----------
def stamp_xlsx(wb, key):
    """Stamp every worksheet's print header+footer and set workbook metadata.
    Header right = 'Version X · date'; footer left = title·v·date, right = Page P of N.
    Renders in the LibreOffice PDF export and in Excel's Page Layout view."""
    for ws in wb.worksheets:
        ws.oddHeader.right.text = cover_line(key)
        ws.oddFooter.left.text = footer_text(key)
        ws.oddFooter.right.text = "Page &P of &N"
    try:
        wb.properties.title = title(key)
        wb.properties.version = ver(key)
        wb.properties.keywords = footer_text(key)
    except Exception:
        pass
