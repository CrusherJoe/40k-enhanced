#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate the LSO Knights Runbook (Word) from tools/lso_data.py.
   PYTHONPATH=tools python3 tools/gen_lso_runbook_docx.py
Output: docs/Reports & Plans/LSO-Runbook.docx
Covers the LIST DECISION (2 Castellan/1 Lancer vs 1 Castellan/2 Lancer) + per-archetype
battle plans, each annotated with which list is better for that matchup.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import lso_data as D
import doc_versions as V

DOCKEY = "lso-runbook"
OUT = V.out_path(DOCKEY)

NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x1F, 0x5C, 0x99)   # List A
ORANGE = RGBColor(0xB0, 0x5A, 0x00)  # List B
GREY = RGBColor(0x60, 0x60, 0x60)
VERDICT_RGB = {
    "FAVOURABLE": RGBColor(0x2E, 0x7D, 0x32), "COIN-FLIP": RGBColor(0xB8, 0x86, 0x00),
    "EXPECT-A-LOSS (winnable)": RGBColor(0xC0, 0x60, 0x00), "UNFAVOURABLE": RGBColor(0xC0, 0x39, 0x00),
    "HARD-LOSS": RGBColor(0xB0, 0x20, 0x00), "AUTO-LOSS": RGBColor(0x99, 0x00, 0x00), "PRELIM": GREY,
}


def lean_color(lean):
    k = lean.split()[0]
    return BLUE if k == "A" else ORANGE if k == "B" else GREY


def vcol(v):
    for k in (v, v.split(" (")[0]):
        if k in VERDICT_RGB:
            return VERDICT_RGB[k]
    return NAVY


def h(doc, text, size, color=NAVY, before=8, after=4, bold=True):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color
    return p


def kv(doc, k, v, kcolor=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(k + ": "); r.bold = True
    if kcolor is not None:
        r.font.color.rgb = kcolor
    p.add_run(v)


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def list_block(doc, name, identity, units, color):
    h(doc, name, 12, color=color, before=6)
    doc.add_paragraph(identity).italic = True
    for uname, cnt, wg, pts, role in units:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{cnt}x {uname} ({pts}) — ").bold = True
        p.add_run(role)


def main():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10)

    # cover
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("LSO RUNBOOK — IMPERIAL KNIGHTS"); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run(f"List Decision + Per-Archetype Battle Plans\n{D.EVENT}\nGenerated {D.GENERATED}").italic = True
    vp = doc.add_paragraph(); vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = vp.add_run(V.cover_line(DOCKEY)); vr.bold = True; vr.font.size = Pt(11); vr.font.color.rgb = NAVY

    # ---- THE DECISION ----
    h(doc, "THE DECISION — which list to take", 16)
    list_block(doc, D.LIST_A_NAME, D.LIST_A_IDENTITY, D.LIST_A_UNITS, BLUE)
    list_block(doc, D.LIST_B_NAME, D.LIST_B_IDENTITY, D.LIST_B_UNITS, ORANGE)
    h(doc, "The trade (verified 11E sim)", 12, color=RGBColor(0, 0, 0), before=8)
    for name, kind, txt in D.SIM_DELTA:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{name} [{kind}]: ").bold = True; p.add_run(txt)
    h(doc, "Prevalence-weighted tally (n=75 top-finishing lists)", 12, color=RGBColor(0, 0, 0), before=8)
    for bucket, items in D.DECISION_TALLY.items():
        col = BLUE if "LIST A" in bucket else ORANGE if "LIST B" in bucket else GREY
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        rr = p.add_run(bucket + ": "); rr.bold = True; rr.font.color.rgb = col
        p.add_run(", ".join(items))
    h(doc, "Recommendation & reasoning", 13, before=8)
    for para in D.DECISION.split("\n\n"):
        doc.add_paragraph(para)

    # ---- THE META ----
    doc.add_page_break()
    h(doc, "THE META — what's winning (n=75)", 15)
    doc.add_paragraph(D.OBSERVED_META_NOTE).italic = True
    tbl = doc.add_table(rows=1, cols=4); tbl.style = "Light Grid Accent 1"; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, x in enumerate(("Faction / archetype", "# top", "Character", "Favours")):
        tbl.rows[0].cells[i].paragraphs[0].add_run(x).bold = True
    for name, cnt, char, fav, note in D.META:
        c = tbl.add_row().cells
        c[0].text = name; c[1].text = str(cnt); c[2].text = char
        rr = c[3].paragraphs[0].add_run(fav); rr.bold = True; rr.font.color.rgb = lean_color(fav)
    for col, w in zip(tbl.columns, (3.0, 0.6, 2.4, 1.0)):
        for cell in col.cells: cell.width = Inches(w)

    # ---- RULES + PROFILES ----
    h(doc, "RULES CHEAT-SHEET", 15)
    for name, desc in D.RULES:
        p = doc.add_paragraph(style="List Bullet"); p.add_run(name + " — ").bold = True; p.add_run(desc)
    h(doc, "Code Chivalric — the Oath (pick 1 Deed + 1 Quality per game)", 12, color=RGBColor(0, 0, 0), before=8)
    dp = doc.add_paragraph(); dp.add_run("Deeds").bold = True
    dp.add_run(" (complete once → Honoured +2/3 CP):")
    for n, d in D.CODE_CHIVALRIC["deeds"]:
        p = doc.add_paragraph(style="List Bullet"); p.add_run(n + " — ").bold = True; p.add_run(d)
    qp = doc.add_paragraph(); qp.add_run("Qualities").bold = True
    qp.add_run(" (army-wide, all game):")
    for n, d in D.CODE_CHIVALRIC["qualities"]:
        p = doc.add_paragraph(style="List Bullet"); p.add_run(n + " — ").bold = True; p.add_run(d)
    h(doc, "Realistic record", 12, color=RGBColor(0, 0, 0), before=8)
    doc.add_paragraph(D.RECORD_NOTE)
    h(doc, "Mindset", 12, color=RGBColor(0, 0, 0), before=6)
    doc.add_paragraph(D.MINDSET)

    # ---- MATCHUP INDEX ----
    doc.add_page_break()
    h(doc, "MATCHUP INDEX", 15)
    tbl = doc.add_table(rows=1, cols=4); tbl.style = "Light Grid Accent 1"; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, x in enumerate(("Archetype", "Verdict", "Better list", "Deciding factor")):
        tbl.rows[0].cells[i].paragraphs[0].add_run(x).bold = True
    for m in D.MATCHUPS:
        lean, _ = D.MATCHUP_LEANS[m["key"]]
        c = tbl.add_row().cells
        c[0].text = f"{m['faction']} — {m['archetype']}"
        rr = c[1].paragraphs[0].add_run(m["verdict"]); rr.bold = True; rr.font.color.rgb = vcol(m["verdict"])
        lr = c[2].paragraphs[0].add_run(lean); lr.bold = True; lr.font.color.rgb = lean_color(lean)
        c[3].text = m["deciding"]
    for col, w in zip(tbl.columns, (2.3, 1.3, 0.9, 3.0)):
        for cell in col.cells: cell.width = Inches(w)

    # ---- BATTLE PLANS ----
    doc.add_page_break()
    h(doc, "BATTLE PLANS", 16)
    for m in D.MATCHUPS:
        lean, why = D.MATCHUP_LEANS[m["key"]]
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.keep_with_next = True
        r = p.add_run(f"{m['faction']} — {m['archetype']}"); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY
        vp = doc.add_paragraph(); vp.paragraph_format.space_after = Pt(2)
        vr = vp.add_run(m["verdict"]); vr.bold = True; vr.font.color.rgb = vcol(m["verdict"])
        vp.add_run(f"   ·   Prevalence: {m['prev']}   ·   They run: {m['disp']}").italic = True
        lp = doc.add_paragraph(); lp.paragraph_format.space_after = Pt(2)
        lr = lp.add_run(f"Better list: {lean}"); lr.bold = True; lr.font.color.rgb = lean_color(lean)
        lp.add_run(f" — {why}")
        kv(doc, "Deciding factor", m["deciding"])
        kdp = doc.add_paragraph(); kdp.paragraph_format.space_after = Pt(2)
        kdp.add_run("The heist:").bold = True
        bullets(doc, m["heist"])
        kv(doc, "Kill-priority", m["kill_priority"])
        kv(doc, "Deploy / positioning", m["deploy"])
        cc = D.CODE_CHIVALRIC["picks"].get(m["key"])
        if cc:
            kv(doc, "Code Chivalric (Deed | Quality)", f"{cc[0]}  |  {cc[1]} — {cc[2]}")

    # ---- VERIFIED PROFILES appendix ----
    doc.add_page_break()
    h(doc, "APPENDIX — VERIFIED 11E PROFILES (the numbers behind the sim)", 14)
    for wname, prof, note in D.VERIFIED_PROFILES:
        if wname.strip().startswith("---"):
            h(doc, "Key enemy anti-Knight weapons", 12, color=RGBColor(0, 0, 0), before=8); continue
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{wname.strip()}: ").bold = True
        p.add_run(prof + (f"  — {note}" if note else ""))

    V.stamp_docx(doc, DOCKEY)
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print("wrote", OUT, V.cover_line(DOCKEY), "-", len(D.MATCHUPS), "battle plans")


if __name__ == "__main__":
    main()
