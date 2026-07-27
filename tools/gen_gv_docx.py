#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Great Value (Imperial Fists) LSO Runbook (Word) from tools/gv_data.py + mc_gv_sim.py.
   PYTHONPATH=tools:src python3 tools/gen_gv_docx.py
Output: docs/Reports & Plans/GV-LSO-Runbook.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import gv_data as D
import doc_versions as V

DOCKEY = "gv-runbook"
OUT = V.out_path(DOCKEY)
GOLD = RGBColor(0x8A, 0x6D, 0x00)   # Imperial Fists yellow-ish
NAVY = RGBColor(0x1F, 0x38, 0x64)
VERD = {"FAVOURABLE": RGBColor(0x2E, 0x7D, 0x32), "COIN-FLIP": RGBColor(0xB8, 0x86, 0x00),
        "HARD": RGBColor(0xB0, 0x20, 0x00)}


def vcol(v):
    for k in VERD:
        if v.upper().startswith(k[:4]):
            return VERD[k]
    return NAVY


def h(doc, text, size, color=GOLD, before=8, after=4):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    return p


def kv(doc, k, v):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    p.add_run(k + ": ").bold = True; p.add_run(v)


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def main():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10)
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("GREAT VALUE — LSO RUNBOOK"); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = GOLD
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run(f"{D.LIST_NAME}\n{D.EVENT}\nGenerated {D.GENERATED}").italic = True
    vp = doc.add_paragraph(); vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = vp.add_run(V.cover_line(DOCKEY)); vr.bold = True; vr.font.size = Pt(11); vr.font.color.rgb = GOLD

    h(doc, "THE LIST", 15)
    kv(doc, "Detachments", D.DETACHMENTS)
    kv(doc, "Disposition", D.DISPOSITION)
    kv(doc, "Points", D.LIST_TOTAL)
    doc.add_paragraph(D.IDENTITY).italic = True
    for name, cnt, role, pts in D.UNITS:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{cnt}x {name} ({pts}) — ").bold = True; p.add_run(role)

    h(doc, "THE TAPESTRY — how it wins", 15)
    for name, desc in D.RULES:
        p = doc.add_paragraph(style="List Bullet"); p.add_run(name + " — ").bold = True; p.add_run(desc)
    h(doc, "Mindset", 12, color=RGBColor(0, 0, 0), before=6)
    doc.add_paragraph(D.MINDSET)
    h(doc, "Realistic record", 12, color=RGBColor(0, 0, 0), before=6)
    doc.add_paragraph(D.RECORD_NOTE)
    for band, lists in D.BANDS.items():
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        p.add_run(band + ": ").bold = True; p.add_run(", ".join(lists))

    # sim verification
    doc.add_page_break()
    h(doc, "SIM VERIFICATION (data-driven)", 15)
    doc.add_paragraph("800 games/archetype. GV's Oath-boosted convergence output computed from real 11e "
                      "profiles (data/bsdata via wh.mathhammer); enemy pressure + objectives meta-calibrated. "
                      "tools/mc_gv_sim.py.").italic = True
    try:
        import mc_gv_sim
        rows = mc_gv_sim.results(800)
    except Exception as e:
        rows = []
        doc.add_paragraph(f"(sim unavailable: {e})")
    if rows:
        tbl = doc.add_table(rows=1, cols=6); tbl.style = "Light Grid Accent 1"; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, x in enumerate(("Archetype", "Prev.", "Verdict", "Oath dmg", "Brick", "GV win%")):
            tbl.rows[0].cells[i].paragraphs[0].add_run(x).bold = True
        for x in rows:
            c = tbl.add_row().cells
            c[0].text = x["archetype"]; c[1].text = str(x["prev"]); c[2].text = x["verdict"]
            c[3].text = str(x["oath"]); c[4].text = str(x["brick"])
            rr = c[5].paragraphs[0].add_run(f"{x['win']}%"); rr.bold = True
        for col, wd in zip(tbl.columns, (2.6, 0.7, 1.0, 1.0, 0.8, 0.9)):
            for cell in col.cells:
                cell.width = Inches(wd)
        prev = sum(x["prev"] for x in rows); wr = sum(x["prev"] * x["win"] for x in rows) / prev
        doc.add_paragraph(f"Prevalence-weighted GV win rate across the top meta ≈ {wr:.0f}% — a strong all-comers "
                          f"list. (Indicative, not gospel; the coin-flips are the matchups to practice.)").italic = True

    # battle plans
    doc.add_page_break()
    h(doc, "BATTLE PLANS", 16)
    for m in D.MATCHUPS:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.keep_with_next = True
        r = p.add_run(f"{m['faction']} — {m['archetype']}"); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = GOLD
        vp = doc.add_paragraph(); vp.paragraph_format.space_after = Pt(2)
        vr = vp.add_run(m["verdict"]); vr.bold = True; vr.font.color.rgb = vcol(m["verdict"])
        vp.add_run(f"   ·   Prevalence: {m['prev']}").italic = True
        kv(doc, "Deciding factor", m["deciding"])
        gp = doc.add_paragraph(); gp.paragraph_format.space_after = Pt(2); gp.add_run("Game plan:").bold = True
        bullets(doc, m["plan"])
        kv(doc, "Watch for", m["watch"])

    doc.add_page_break()
    h(doc, "APPENDIX — VERIFIED PROFILES (the numbers behind the sim)", 13)
    for wn, prof, note in D.VERIFIED_PROFILES:
        p = doc.add_paragraph(style="List Bullet"); p.add_run(wn + ": ").bold = True
        p.add_run(prof + (f"  — {note}" if note else ""))

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    V.stamp_docx(doc, DOCKEY)
    doc.save(OUT)
    print("wrote", OUT, V.cover_line(DOCKEY), "-", len(D.MATCHUPS), "battle plans")


if __name__ == "__main__":
    main()
