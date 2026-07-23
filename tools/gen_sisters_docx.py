#!/usr/bin/env python3
"""Build the Adepta Sororitas all-comers tournament battle plan as a Word doc.
Pulls the live roster/points from the army builder and the deployment layouts
from data/layouts/disruption.yaml, so the doc stays in sync with the data.

    PYTHONPATH=src python3 tools/gen_sisters_docx.py
-> docs/Sisters-Battle-Plan.docx
"""
import os
os.environ.setdefault("WH_FACTION", "sisters")
import yaml
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from wh import army as army_mod, data

doc = Document()
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
CRIMSON = RGBColor(0x7A, 0x0E, 0x2A)   # Sororitas red
STEEL = RGBColor(0x2B, 0x3A, 0x4A)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = CRIMSON if level == 1 else STEEL
    return p


def para(*segs, style=None, after=4, italic=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    for s in segs:
        if isinstance(s, tuple):
            run = p.add_run(s[0]); run.bold = s[1]
        else:
            run = p.add_run(s)
        run.italic = italic
    return p


def bullet(*segs, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(2)
    for s in segs:
        if isinstance(s, tuple):
            run = p.add_run(s[0]); run.bold = s[1]
        else:
            p.add_run(s)
    return p


# ---------------- title ----------------
t = doc.add_heading('ADEPTA SORORITAS', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in t.runs:
    r.font.color.rgb = CRIMSON
sub = para(('All-Comers Tournament Battle Plan  ·  Champions of Faith + Sacred Champions  ·  LOCK: DISRUPTION', True))
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
para('Derived from the current 11E meta (5 factions / 7 archetypes) and validated with the '
     'mathhammer engine. Points are exact (MFM). All rules verified against 11E sources.',
     italic=True, after=10)

# ---------------- the list ----------------
h('1 · The Army List', 1)
a = army_mod.build_file('examples/best-sisters-allcomers.yaml')
para((f'{a.name}  —  {a.points}/{a.points_limit} pts', True))
para(('Detachments: Champions of Faith (2DP, Disruption) + Sacred Champions (1DP, Take-and-Hold). '
      'Rules stack army-wide.', False), after=6)
for l in a.lines:
    label = l.datasheet + (f' [{l.models}]' if l.models else '')
    extra = ''
    if l.wargear:
        extra += '  ·  ' + ', '.join(l.wargear)
    if l.enhancement:
        extra += f'  ·  {l.enhancement}'
    bullet((f'{label}', True), (f'  {l.total} pts{extra}', False))

h('Attachments (Leader + Support per unit, Core 19.01)', 2)
bullet(('Vahl → Paragon Warsuits', True), '  — re-roll Hits & Wounds (her unit only) + Righteous +1 BS. The melta hammer.')
bullet(('Palatine [Triptych of Judgement] + Imagifier → Multi-melta Retributor', True),
       '  — Lethal Hits + ignores cover + the whole unit becomes Sv 2+/4++. Durable 2nd anchor-killer.')
bullet(('Canoness [Sanctified Amulet] + Dogmata → Celestian Sacresants', True),
       '  — anti-Deep-Strike bubble + +1 OC on a 4++ / Holy-Quest (+1 BS/WS) mid-board holder.')

# ---------------- how it works ----------------
h('2 · How the army works — SOFTEN, then DELETE', 1)
para(('Naked Sisters melta into cover is useless (~4.6 dmg vs a Land Raider). The list wins by '
      'softening a target first, army-wide, then shooting it with buffed melta (~17.8 — one-shots it). '
      'The enablers ARE the anti-tank.', False), after=6)
for step, txt in [
    ('1. Strip cover', 'Immolator "Purge & Cleanse" (auto-hits with Immolation Flamers) → target loses Benefit of Cover for the whole army. ×2 targets/turn.'),
    ('2. Improve AP', 'Castigator "Rites of Castigation" → +1 AP for every Sororitas gun into that target. ×1 target/turn.'),
    ('3. Buff the shooter', 'Vahl re-roll Hits+Wounds + Righteous Purpose +1 BS on the Paragons = near-auto-hitting, re-rolling melta.'),
    ('4. Cover-immune backup', 'Palatine + Triptych on the MM Retributor → that unit ignores cover on its own, freeing an Immolator strip for another target.'),
]:
    bullet((step + ': ', True), txt)
para(('Per-turn sequence: ', True),
     'Immolators shoot first (strip cover) → Castigator marks the key target (+1 AP) → Vahl+Paragons '
     'delete anchor A → MM-Retributor + a melta Dominion delete anchor B → flamers / heavy bolters / '
     'jump clear chaff and screen. Spend a Miracle die for a clutch melta WOUND on a high-T target.', after=8)

# ---------------- cheat sheet ----------------
h('3 · Mathhammer cheat-sheet (full hit→wound→save chain)', 1)
h('Anti-tank (buffed = softened target, half range)', 2)
for row in [
    ('Vahl+Paragons (6 MM)', 'Land Raider T12/2+ : 17.8 (kills 16W)  ·  Monolith T13/2+ : 17.8 (81%)  ·  Deathwing 2+/4++ : 14.3 (~3.5 Termies)'),
    ("C'tan T11/3+/4++, −1 Damage", '7.3 of 16W even fully buffed → DON\'T bother, play the mission'),
]:
    bullet((row[0] + ' — ', True), row[1])
h('Anti-horde vs 20 Ork Boyz (T5 W1 5+) — dead/turn', 2)
para('Castigator 14.6  +  2× Immolation Flamers 7.8  +  flamer Dominion 4.7  +  heavy-bolter Retributor 6.7  '
     '+  Zephyrim 3.1  ≈  37 dead Boyz/turn (break blobs off objectives).', after=6)
h('Durability — how fast my girls die (models lost/turn)', 2)
bullet(('Cheap bodies evaporate to VOLUME: ', True),
       'one Boyz mob or a 10-pyreblaster wall WIPES a 10-Sister squad (even the 4++). '
       'Battle Sisters are screens / traders / backfield OC — never front-line, never inside 12" of flamers.')
bullet(('Imagifier 2+/4++ brick ≈ 1.7× tougher vs anti-tank sniping', True),
       ' (death rays 2.8→1.7) — its job — but volume still swamps it. Keep it BACK, sniping anchors.')
bullet(('The hammer delivers: ', True),
       'Paragons take only 6.0 from a full Monolith\'s death rays (survive), 2.7 from a Repulsor.')
bullet(('Transports are durable delivery: ', True),
       'an Immolator shrugs 6 kopta rokkits (4.4) and survives 3 gauss destructors — not glass, but extract value early.')

# ---------------- disposition lock ----------------
h('4 · Disposition — LOCK DISRUPTION', 1)
para(('In a tournament you lock ONE disposition for the whole event. Lock DISRUPTION: it is favoured '
      'against 4 of the 6 boogeymen and suits this mobile, mission-first list. Your MISSION is then set '
      'by the OPPONENT\'s disposition:', False), after=6)
for opp, mis in [
    ('opponent Purge the Foe', 'you play DELAYING ACTION (2 VP/kill + hold non-home)'),
    ('opponent Reconnaissance', 'you play SMOKE AND MIRRORS (Decoy objectives, esp. in their half; 10 VP if 4+ decoyed)'),
    ('opponent Disruption', 'you play OUTMANOEUVRE (escalating per-objective; 10 VP for THEIR home)'),
    ('opponent Take-and-Hold', 'you play DEATH TRAP (Booby-Trap terrain + kill enemies in it + hold non-home)'),
]:
    bullet((opp + ' → ', True), mis)

# ---------------- deployment layouts ----------------
h('5 · Deployment layouts (Disruption matchups, from the Event Companion)', 1)
para(('Universal on every layout: 6 objectives (2 home deep in each zone, 2 central in No Man\'s Land, '
      '2 expansion on the flanks), 16 obscuring terrain areas, centre packed (no turn-1 cross-board LOS). '
      'Default your plan to LAYOUT A; B/C notes below.', False), after=6)
lay = yaml.safe_load(open('data/layouts/disruption.yaml'))
for m in lay['matchups']:
    if m['vs'] == 'priority-assets':
        continue
    para((f"vs {m['vs'].replace('-', ' ').title()}  (you: {m['your_mission']})", True), after=2)
    for L in ('A', 'B', 'C'):
        d = m['layouts'][L]
        bullet((f'Layout {L}: ', True), f"{d['deployment']} — {d['note']}", level=1)

para(('Standing deployment for this list (all layouts): ', True),
     'castle the Vahl+Paragon hammer + the 2+/4++ MM-Retributor brick central-rear with LOS to the mid; '
     'Immolators + Dominions mid-forward (hull-down T1); Sacresants+Canoness hold centre with the '
     'anti-Deep-Strike bubble; Battle Sisters screen + sit home/expansion; Zephyrim/Seraphim flanks or '
     'Reserve. Diagonal (A) = long approach, angle the hammer across the mid. Horizontal/Dawn-of-War (C) '
     '= close, screen harder and value the bubble. Vertical (B) = cover the width with the jump units.',
     after=8)

# ---------------- per-archetype plans ----------------
h('6 · Per-archetype battle plans (locked Disruption)', 1)
plans = [
    ('Salamanders flamer-brick', 'Recon / Purge', 'Smoke and Mirrors (vs Recon) or Delaying Action (vs Purge)',
     'Soften + delete the 2 Land Raiders (Immolator strip → Castigator +1 AP → Vahl+Paragons one, MM-Retributor + melta Dominion the other), then Aggressors/Bladeguard.',
     'Do NOT walk bodies inside 12" of the pyreblasters (auto-wipe). Score kills as tempo, Decoy from range, counter-charge stragglers with Zephyrim.'),
    ('Ork Kult of Speed (dakka)', 'Disruption / Purge', 'Outmanoeuvre (vs Disruption) or Delaying Action (vs Purge)',
     'Melta Kill Rigs + Wazdakka; Deffkoptas/Flash Gitz with Castigator Blast + heavy bolters + flamers (AP-1/-2 volume, not pure AP-4).',
     'They are faster — anchor on objectives and make them come; win the grind while they trade inefficiently. Don\'t over-expose transports (multiple AT units bracket them over 2 turns).'),
    ('Ork Green Tide (100 Boyz)', 'Take-and-Hold', 'Death Trap',
     'The flamer game — Castigator + 2 Immolation-Flamer Immolators + flamer Dominion + heavy bolters ≈ 37 dead Boyz/turn. Kill WHOLE units to break OC. Save melta for Ghazghkull / wagons.',
     'Boyz cluster in terrain → Booby-Trap those areas + clear the blobs there for double VP, and dodge the OC race. NEVER get charged — screen, Overwatch flamers, fall back + shoot, hold with Sacresants (4++).'),
    ('Necron Monoliths (3× T13, OC8, slow)', 'Disruption', 'Outmanoeuvre',
     'Do NOT dogpile a Monolith. Kill the Lokhust / Silent-King support / Ophydians; deny their scoring.',
     'They are M8" and only three models — out-maneuver them, grab the objectives they can\'t reach, race their home (10 VP). Screen Ophydian Deep Strike with the Amulet bubble.'),
    ("Necron C'tan (4× T11, −1 Damage)", 'Purge / Recon', 'Delaying Action (vs Purge) or Smoke and Mirrors (vs Recon)',
     'IGNORE the C\'tan (un-killable — 7.3 even fully buffed). Melta/kill the killable enablers: Lychguard, Lokhust, Reanimator, Plasmancer, characters (2 VP each on Delaying Action).',
     'Their Purge list farms your dead units — feed them NOTHING. Keep fragile units out of C\'tan + Drain-Life (6" mortal) range; hold central + expansion while the slow C\'tan sit.'),
    ('Dark Angels Deathwing (2+/4++ bricks, OC1)', 'Take-and-Hold', 'Death Trap',
     'Melta works on the W4 Terminators (multi-damage past the 4++) — chip a softened brick, but don\'t over-commit. Kill the Sternguard / Eradicators / Repulsor / Scouts (the scoring pieces).',
     'The bricks are near-immortal but OC1 — OUT-OC them: flood the non-home objectives (4 VP) with bodies. SCREEN the Deep Strike hard with the Sanctified Amulet — deny the teleport and the bricks walk slowly into your guns.'),
]
for name, their_disp, mission, targets, plan in plans:
    h(name, 2)
    para(('They lock: ', True), their_disp, ('   →   You play: ', True), mission, after=2)
    bullet(('Targets: ', True), targets)
    bullet(('Plan: ', True), plan)

# ---------------- footer ----------------
para(('Sources: ', True),
     'list examples/best-sisters-allcomers.yaml; mechanics docs/sisters-mechanics.md; mathhammer + '
     'durability docs/sisters-battle-plan.md; matchups docs/sisters-matchup-plans.md; layouts '
     'data/layouts/disruption.yaml. All points MFM; all rules verified 11E.', italic=True, after=0)

out = 'docs/Sisters-Battle-Plan.docx'
doc.save(out)
print(f'wrote {out}  ({a.points}/{a.points_limit} pts, {len(a.lines)} units, {len(plans)} matchup plans)')
