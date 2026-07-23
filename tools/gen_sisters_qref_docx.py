#!/usr/bin/env python3
"""Build a printable, tabletop per-round QUICK-REFERENCE card for the Sisters
all-comers list (locked Disruption). Compact, one/two pages.

    PYTHONPATH=src python3 tools/gen_sisters_qref_docx.py
-> docs/Sisters-Quick-Reference.docx
"""
import os
os.environ.setdefault("WH_FACTION", "sisters")
import yaml
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
for s in doc.sections:                      # tight margins for printing
    s.top_margin = s.bottom_margin = Inches(0.5)
    s.left_margin = s.right_margin = Inches(0.55)
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(9.5)
CRIMSON = RGBColor(0x7A, 0x0E, 0x2A)
STEEL = RGBColor(0x2B, 0x3A, 0x4A)


def band(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = CRIMSON
    return p


def line(*segs, after=1, size=9.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    for s in segs:
        r = p.add_run(s[0] if isinstance(s, tuple) else s)
        r.bold = isinstance(s, tuple) and s[1]
        r.font.size = Pt(size)
    return p


def tick(*segs):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.2)
    for s in segs:
        r = p.add_run(s[0] if isinstance(s, tuple) else s)
        r.bold = isinstance(s, tuple) and s[1]
        r.font.size = Pt(9.5)
    return p


# ---- title ----
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('SISTERS — TABLETOP QUICK REFERENCE')
r.bold = True; r.font.size = Pt(15); r.font.color.rgb = CRIMSON
t.paragraph_format.space_after = Pt(0)
st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = st.add_run('LOCK: DISRUPTION   ·   1985/2000   ·   Champions of Faith + Sacred Champions')
sr.bold = True; sr.font.size = Pt(9.5); sr.font.color.rgb = STEEL
st.paragraph_format.space_after = Pt(3)

# ---- mission by opponent (table, from the layout data) ----
band('MISSION — set by OPPONENT’s disposition')
lay = yaml.safe_load(open('data/layouts/disruption.yaml'))
mrow = {m['vs']: m for m in lay['matchups']}
pairs = [('Purge the Foe', 'purge-the-foe'), ('Reconnaissance', 'reconnaissance'),
         ('Disruption', 'disruption'), ('Take-and-Hold', 'take-and-hold')]
tbl = doc.add_table(rows=1, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for i, (label, key) in enumerate(pairs):
    hc = hdr[i].paragraphs[0]; hr = hc.add_run('opp ' + label); hr.bold = True; hr.font.size = Pt(8.5)
row = tbl.add_row().cells
for i, (label, key) in enumerate(pairs):
    rc = row[i].paragraphs[0]; rr = rc.add_run(mrow[key]['your_mission']); rr.font.size = Pt(9); rr.bold = True
line('', after=2)

# ---- deploy ----
band('DEPLOY  (default Layout A/B = diagonal, long approach · Layout C = close → screen harder)')
tick(('Castle rear-centre: ', True), 'Vahl+Paragons hammer + the 2+/4++ MM-Retributor brick, LOS to the mid.')
tick(('Mid-forward (hull-down T1): ', True), '2 Immolators (flamers) + 2 Dominions (melta / flamer split-halves).')
tick(('Hold centre: ', True), 'Sacresants + Canoness (ANTI-DEEP-STRIKE bubble) + Dogmata (+1 OC).')
tick(('Screen / OC / actions: ', True), 'Battle Sisters on home+expansion; Zephyrim & Seraphim flanks or Reserve.')

# ---- softening sequence ----
band('EVERY SHOOTING PHASE — SOFTEN, then DELETE')
line(('  1. ', True), 'IMMOLATORS shoot first (auto-hit) → strip Benefit of Cover on 2 priority targets.')
line(('  2. ', True), 'CASTIGATOR shoots → +1 AP (Rites of Castigation) on the single key target.')
line(('  3. ', True), 'VAHL+PARAGONS → delete anchor A  (re-roll H&W, BS2+, half-range Melta 2).')
line(('  4. ', True), 'MM-RETRIBUTOR + a melta DOMINION → anchor B.')
line(('  5. ', True), 'Flamers / heavy bolters / jump → clear chaff + screen. Miracle die = a clutch melta WOUND.')

# ---- target priority ----
band('TARGET PRIORITY')
tick(('MELTA (only when softened): ', True), 'Land Raiders, Kill Rigs, ONE Monolith, Repulsor, W4 Terminators.')
tick(('FLAMER / BLAST: ', True), 'hordes & chaff (Boyz, Deffkoptas) — Castigator + Immolation flamers + heavy bolters (~37 Boyz/turn).')
tick(('IGNORE (un-killable): ', True), 'C’tan, extra Monoliths → play the MISSION, don’t dogpile (melta is 3× overkill on one target).')
tick(('KILL THE ENABLERS: ', True), 'enemy support characters, Reanimator, Lokhust, Lychguard, Sternguard/Eradicators, Scouts.')

# ---- per round ----
band('PER ROUND')
tick(('T1: ', True), 'Advance to open lanes; strip + kill anchor A; start actions (Decoy / Booby-Trap); grab a central objective; screen the Deep Strike.')
tick(('T2–3 (peak VP): ', True), 'Kill anchor B; score primary HARD; contest expansion; jump units for enemy home / actions.')
tick(('T4–5: ', True), 'Hold what you have; deny their scoring; trade nothing; close on primary + secondaries.')

# ---- don't ----
band('DON’T  (T3 W1 — girls die a lot)')
line('✖ Sit within 12" of flamers (auto-wipes a squad).   ✖ Let a horde charge you.   '
     '✖ Over-melta the un-killable.   ✖ Over-expose transports.   ✖ Dogpile one anchor.', after=0, size=9)

out = 'docs/Sisters-Quick-Reference.docx'
doc.save(out)
print(f'wrote {out}')
