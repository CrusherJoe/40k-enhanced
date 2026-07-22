#!/usr/bin/env python3
"""Build the Lancer/Purge battle plan as a formatted Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# --- base styles ---
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)

NAVY = RGBColor(0x1F, 0x33, 0x55)
RED = RGBColor(0x8B, 0x1A, 0x1A)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY if level > 1 else RED
    return p


def para(*segments, style=None, space_after=4):
    """segments: strings or (text, bold) tuples."""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    for seg in segments:
        if isinstance(seg, tuple):
            run = p.add_run(seg[0]); run.bold = seg[1]
        else:
            p.add_run(seg)
    return p


def bullet(*segments, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(2)
    for seg in segments:
        if isinstance(seg, tuple):
            run = p.add_run(seg[0]); run.bold = seg[1]
        else:
            p.add_run(seg)
    return p


def phase(name, *segments):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(name + ': '); r.bold = True; r.font.color.rgb = NAVY
    for seg in segments:
        if isinstance(seg, tuple):
            run = p.add_run(seg[0]); run.bold = seg[1]
        else:
            p.add_run(seg)
    return p


# ============================ TITLE ============================
title = doc.add_heading('', level=0)
r = title.add_run('Imperial Knights — Battle Plan'); r.font.size = Pt(24); r.font.color.rgb = RED
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = sub.add_run('Valourstrike Lance + Dominus Foebreakers  ·  Force Disposition: Purge the Foe  ·  2000 pts')
r.italic = True; r.font.size = Pt(11); r.font.color.rgb = NAVY
doc.add_paragraph()

# ============================ THE LIST ============================
h('The List (1995 / 2000)', 1)
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Light Grid Accent 1'
hdr = tbl.rows[0].cells
for i, t in enumerate(['Unit', 'Wargear / Enhancement', 'Job']):
    hdr[i].paragraphs[0].add_run(t).bold = True
rows = [
    ('Knight Castellan', 'Archeotech Autoloaders', 'Main guns; re-rolls its casino dice (no more whiffed volcano lance)'),
    ('Knight Castellan', 'Blessed Plate (T13)', 'Main guns; the durable anchor'),
    ('Knight Crusader', 'Rapid-fire battle cannon + Judicant’s Helm', 'Flex guns; feeds IGNORES COVER to a Castellan each Shooting phase'),
    ('Cerastus Knight Lancer', '—', 'Distraction + melee alpha; 4+ invuln vs everything, 14" move, S20 lance'),
    ('Armiger Warglaive', '—', 'Grabber/finisher; gets advance-and-charge from the Lancer’s Bondsman'),
    ('Navigator (Agents ally)', '— (75 pts)', 'Home-objective anchor; Hidden + Gone to Ground + anti-Deep-Strike bubble'),
]
for a, b, c in rows:
    cells = tbl.add_row().cells
    cells[0].paragraphs[0].add_run(a).bold = True
    cells[1].paragraphs[0].add_run(b)
    cells[2].paragraphs[0].add_run(c)
doc.add_paragraph()

# ============================ WIN CONDITION ============================
h('How this list wins Purge the Foe', 1)
para('Every one of your five missions pays for the ', ('same two things', True),
     ': ', ('killing units', True), ' (3–5 VP most turns) and ',
     ('holding objectives other than your home', True),
     ' (4 VP/turn from round 2), with a late bonus for taking the ',
     ('opponent’s home or a central objective', True), ' (5–8 VP). You do both natively — the guns kill, '
     'and your big Knights ', ('hold with OC 10 while they shoot', True),
     ' (56–80" range, so they never choose move-or-shoot). The Navigator locks your home so you '
     'bank that VP for free and commit everything else forward.')

h('Two habits that win games', 2)
para(('1. Free mobility (Valourstrike). ', True),
     'The instant ANY Knight Advances, every Knight’s ranged weapons gain [ASSAULT] for the turn. '
     'So you always Advance-and-shoot — move out of hiding, into a lane, or onto an objective, and still fire '
     'the full salvo. Treat the whole gunline as mobile.')
para(('2. Resolve the cover stack every Shooting phase. ', True),
     'Cover is a −1 to hit and it is the DEFAULT (16 obscuring terrain areas). Beat it, in priority order per shot:')
bullet(('Judicant', True), ' → the Castellan the Crusader feeds ignores cover entirely. Point it at your best covered target.')
bullet(('Dominus (+1 vs a target in a terrain area)', True), ' → cancels cover for a Castellan shooting something standing in terrain.')
bullet(('Plunging Fire', True), ' → a Castellan (TOWERING) shooting a GROUND-level target WITHIN 12" gets +1 BS (offsets cover). Close range only.')
bullet(('Positioning', True), ' → angle for lanes with no intervening terrain. You have the range; you’re picking angles, not closing distance.')

# ============================ DEPLOYMENT ============================
doc.add_page_break()
h('Pre-game & Deployment', 1)
para('At muster, lock ', ('Purge the Foe', True), ' as your disposition. Then, in deployment order:')
bullet(('Navigator first — on your home objective, inside a DENSE terrain feature', True),
       ' so it is Hidden. It never moves and never shoots all game. If the home objective has no dense terrain, '
       'place it behind the nearest dense piece still within 3" of the marker.')
bullet(('Hide the gunline turn 1.', True), ' Deploy the two Castellans and the Crusader behind terrain — you are a '
       'low-model army, so do NOT gift a turn-1 alpha strike clean shots, especially going second. You will emerge '
       'turn 1 with Assault anyway.')
bullet(('Keep the Crusader within 12" of a Castellan', True), ' so Judicant can feed it Ignores Cover next turn.')
bullet(('Lancer on the flank you intend to pressure', True), ', poised to run up a side lane. ',
       ('Warglaive beside it', True), ' (within Bondsman range) so it inherits advance-and-charge.')
bullet(('First turn:', True), ' if you get it, deploy a little more forward and take early lanes. Going second, prioritise hiding.')

# ============================ TURN BY TURN ============================
doc.add_page_break()
h('The Five-Turn Arc (phase by phase)', 1)
para('This is the skeleton every game follows. Distances are Move + Advance (Advance = D6, ~3.5" average) — '
     'measure on the day; treat the numbers as intent. Adapt to the deployment geometry (see the next section).', style=None)

# ---- TURN 1 ----
h('Turn 1 — Emerge and draw first blood', 2)
para(('Goal:', True), ' get out of hiding into firing lanes, take a central objective, kill the scariest thing you can see. '
      'Do NOT overcommit — round 1 is positioning + first blood.')
phase('Command', 'Gain 1 CP. Nothing to battle-shock yet. Note the two central objectives you want (they sit in No Man’s Land terrain).')
phase('Movement', 'Advance the ', ('Lancer first', True), ', ~14 + D6 (≈17.5") up its flank — this switches ',
      ('[ASSAULT] on for the whole army', True), ' and starts the distraction. Then ',
      ('Advance Castellan A ~8 + D6 (≈11.5")', True), ' onto/next to the nearer central objective; ',
      ('Advance Castellan B', True), ' toward the second central or a clean firing lane; the ',
      ('Crusader', True), ' moves up staying within 12" of a Castellan; the ',
      ('Warglaive', True), ' follows the Lancer. The ', ('Navigator does not move', True), '.')
phase('Shooting', 'Resolve the cover stack: ', ('Crusader’s Judicant', True), ' gives one Castellan Ignores Cover for its best covered target; '
      'point the Dominus Castellans at anything standing in terrain (+1). ', ('Delete the enemy’s single most dangerous unit', True),
      ' you can see — that banks the 3 VP for “a unit was destroyed.” The Navigator stays SILENT (shooting breaks Hidden/Gone to Ground).')
phase('Charge', 'Usually none yet — keep the Lancer as a turn-2 threat rather than a wasted turn-1 charge into a screen. If a soft, high-value '
      'unit is in reach, the Lancer or the Warglaive (Bondsman advance-charge) can take it.')
phase('Fight', 'Only if you charged.')
phase('End of turn', 'You should hold at least one central objective (Castellan sitting on it) and have killed one unit. Lancer is upfield.')

# ---- TURN 2 ----
h('Turn 2 — The kill-and-hold turn (primary scoring opens)', 2)
para(('Goal:', True), ' bank your first big primary tick — hold ≥1 non-home objective (4) + a kill (3) — and start dictating the board.')
phase('Command', 'Gain 1 CP. Battle-shock any of your units below half-strength (rare this early). ',
      ('If you are on Punishment (vs Disruption):', True), ' this is when you CONDEMN 1–3 enemy units (ones near objectives '
      'or that killed your stuff) — pick the units you were going to shoot anyway.')
phase('Movement', 'Settle the ', ('Castellans onto both central objectives', True), ' (OC 10 holds them) and into their best lanes — '
      'Advance if you need the reach, you still shoot. Reposition the ', ('Crusader', True), ' to keep a Castellan within 12". Move the ',
      ('Lancer', True), ' into charge range of a high-value target or onto the enemy’s expansion objective; move the ',
      ('Warglaive', True), ' up (advance-and-charge available).')
phase('Shooting', 'Cover stack again. Prioritise: (a) units that threaten your Knights, (b) mission targets (condemned units on Punishment). '
      'Keep the volume on — Archeotech means Castellan A’s casino guns don’t whiff.')
phase('Charge', ('Lancer charges', True), ' its first victim (S20, ', ('+1 to wound on the charge', True),
      ', ~18–30 dmg) — or bodies onto an objective to contest it. ', ('Warglaive advance-charges', True), ' a tank or backfield unit.')
phase('Fight', 'Resolve Lancer/Warglaive combats. Pile in 3", swing, consolidate 3" toward the next objective or target.')
phase('End of turn', '≈7 VP primary banked (kill 3 + hold 4). You hold your home (Navigator) plus a central; the Lancer is in their business.')

# ---- TURN 3 ----
h('Turn 3 — Press the advantage (the swing turn)', 2)
para(('Goal:', True), ' win the kill-race and extend to a second scoring objective. This is usually where the game is decided.')
phase('Command', 'Gain 1 CP. Re-condemn on Punishment. Check battle-shock on anything chewed down.')
phase('Movement', 'Push a Knight toward a SECOND non-home objective (you want “control ≥1 / more than the opponent” every turn). '
      'Aim the ', ('Lancer at the enemy’s home objective', True), ' now — it’s the late-game payout (5–8 VP in most of your missions). '
      'Keep at least one Castellan camped on a central.')
phase('Shooting', 'Trade UP — Meatgrinder and Destroyer’s Wrath pay 4–5 VP for “more enemy units destroyed this turn than friendly last turn.” '
      'Focus-fire to finish units rather than spreading wounds. Full cover stack.')
phase('Charge', 'Lancer charges to clear a path to the enemy home / kill a key scorer; Warglaive finishes a wounded target.')
phase('Fight', 'Resolve; consolidate toward objectives.')
phase('End of turn', 'You should be holding more objectives than the opponent and out-killing them. If ahead, start thinking about how to '
      'close; if behind on bodies, use the Knights’ durability to grind.')

# ---- TURN 4 ----
h('Turn 4 — Convert', 2)
para(('Goal:', True), ' turn board control into locked-in VP; set up the end-of-battle bonus.')
phase('Command', 'CP; condemn/battle-shock as above.')
phase('Movement', 'Consolidate onto the objectives you can actually hold to the end. Move the ', ('Lancer / Warglaive onto or beside the '
      'enemy home objective', True), ' (Meatgrinder/Punishment: controlling it is 5–8 VP; Consecrate: get a unit that just made a kill onto it '
      'so you can consecrate it).')
phase('Shooting', 'Keep killing for the per-turn 3 VP and the “more than last turn” bonus; clear enemy units off contested objectives.')
phase('Charge / Fight', 'Lancer/Warglaive take or hold the enemy home; finish wounded units.')
phase('End of turn', 'Ideally: your home (Navigator) + ≥1 central + pressure on the enemy home. Watch ',
      ('Deadly Demise', True), ' — if a Knight is about to die, face it into clustered enemies (on a 6 it deals mortal wounds within 6").')

# ---- TURN 5 ----
h('Turn 5 — Close it out', 2)
para(('Goal:', True), ' bank the end-of-battle points (they ignore the per-round cap) and deny theirs.')
phase('Command', 'Final CP; last condemn on Punishment.')
phase('Movement', 'Get bodies onto every objective you can score at end of battle — especially a ', ('central objective (Unstoppable Force: 5 VP)',
      True), ' and the ', ('enemy home (Meatgrinder 5 / Punishment 8 / Vanguard-style plays)', True),
      '. The Navigator has held your home all game.')
phase('Shooting / Charge / Fight', 'Kill anything contesting the objectives you need; make the last-turn kills count. Use the Lancer as a '
      'body on an objective if the charge isn’t there.')
phase('End of battle', 'Score the end-of-battle conditions. With the Navigator on home and Knights on the centrals, you close on primary '
      'while having out-killed them all game.')

# ============================ GEOMETRY ============================
doc.add_page_break()
h('Deployment-Geometry Adaptations (A / B / C rotate through these)', 1)
para(('Horizontal / Dawn of War (~18" zones, ~8" No Man’s Land — CLOSE): ', True),
     'you can pressure turn 1–2. Hide behind the mid-line terrain, emerge and shoot early; the Lancer reaches combat by turn 2. '
     'Cuts both ways — respect the enemy’s alpha strike.')
para(('Vertical / Hammer-and-Anvil (~12–14" zones, wide No Man’s Land — FAR): ', True),
     'long approach. Lean on the gunline (you out-range everything) and use Assault’s extra reach to walk up while shooting. The Lancer '
     'needs ~2 turns to connect, so use it to threaten a flank objective rather than an early charge. More time to gunline = good for you.')
para(('Diagonal / Crucible / Tipping-point (corner-to-corner): ', True),
     'castle the gunline in a corner with good sightlines; the centre terrain cluster blocks cross-board line of sight, so take an early '
     'lane on one diagonal. Run the Lancer up the long diagonal at the enemy’s corner. Navigator anchors your corner home.')

# ============================ MATCHUPS ============================
doc.add_page_break()
h('The Five Matchups (your mission vs theirs)', 1)


def matchup(title, they, body_segments):
    h(title, 2)
    p = doc.add_paragraph()
    r = p.add_run('They play: '); r.bold = True; r.font.color.rgb = NAVY
    p.add_run(they)
    para(*body_segments)


matchup('vs Take and Hold → you play Unstoppable Force', 'Immovable Object (sit and hold everything).',
        ['Your edge: you out-shoot a static army and Purge pays you for killing. ',
         ('Kill their holders off the mid-board objectives, then park Knights on them.', True),
         ' End-of-battle bonus: ', ('control a central objective (5 VP)', True), ' — make sure a Knight is sitting central on turn 5.'])
matchup('vs Purge the Foe → Meatgrinder (mirror)', 'Meatgrinder (a kill-race).',
        ['You win it: more, tougher, longer-ranged guns plus the Lancer. ',
         ('Trade up every turn', True), ' (“more enemy units destroyed this turn than friendly last turn” = 5 VP). Late, drive the ',
         ('Lancer onto their home objective (5 VP)', True), '. Protect your Knights’ positioning — don’t feed them into return fire.'])
matchup('vs Disruption → Punishment', 'Delaying Action.',
        [('Special: ', True), 'at the start of each turn you CONDEMN 1–3 enemy units (near objectives / that killed your stuff); you score ',
         ('5 VP whenever a condemned unit leaves the battlefield', True), ' — i.e. when you kill it. Tailor-made for your guns: ',
         ('condemn what you were going to shoot anyway, then delete it.', True),
         ' Hold ≥1 objective (4) + more than them (5); late, take ', ('their home (8 VP)', True), ' with the Lancer/Warglaive.'])
matchup('vs Reconnaissance → Consecrate', 'Triangulation.',
        [('Special: ', True), 'each unit that DESTROYS an enemy becomes a “consecration unit”; at end of turn it consecrates an objective it '
         'is within range of (place a marker). So ', ('kill enemies while your Knights are on/near objectives', True),
         ' → those objectives consecrate themselves (3 VP for 1–2, 6 VP for 3+). Drive a fresh killer onto ',
         ('their home to consecrate it late (5 VP)', True), '. Recon armies are fast/elusive — use the Lancer to hunt their scoring pieces.'])
matchup('vs Priority Assets → Destroyer’s Wrath', 'Vital Link (leans on Objective Actions and operation markers).',
        ['You don’t care about their actions — ', ('kill their action-doers', True), ' (3 VP + the “more than last turn” 4 VP) and ',
         ('out-hold them (6 VP for more objectives than the opponent)', True),
         '. Your OC-10 Knights win the objective count; their fragile action units die to your guns.'])

# ============================ QUICK REFERENCE ============================
doc.add_page_break()
h('Quick-Reference Decision Rules', 1)
bullet(('Advance something every turn', True), ' — it turns on army-wide [ASSAULT], so you always move-and-shoot.')
bullet(('Every Shooting phase, resolve the cover stack first:', True), ' Judicant to the best covered target, Dominus Castellans at in-terrain '
       'targets, note sub-12" ground targets for Plunging Fire.')
bullet(('Target priority:', True), ' kill the thing that most threatens your Knights or scores you the mission (condemned units on Punishment; '
       'units near objectives on Consecrate). Purge pays for the kill regardless.')
bullet(('The Lancer is a THREAT first, a hammer second.', True), ' Position it so the opponent must either divert fire (your gunline shoots '
       'free) or eat an S20 charge. Don’t throw it into a screen it can’t break — angle it at soft, high-value targets and the enemy home.')
bullet(('Hold with Knights, not hope.', True), ' A Castellan on an objective holds it (OC 10) AND shoots — default to parking a Knight on '
       'each central objective you can reach.')
bullet(('Never fire the Navigator.', True), ' Ever. Shooting breaks Hidden / Gone to Ground and opens it to sniping.')
bullet(('Watch Deadly Demise.', True), ' A dying Knight is a bomb — if one is going down, face it into clustered enemies.')

doc.add_paragraph()
foot = doc.add_paragraph()
r = foot.add_run('Generated from the wh project (github.com/CrusherJoe/40k-enhanced) — data, rules and math all verified against the 11th-edition sources.')
r.italic = True; r.font.size = Pt(8); r.font.color.rgb = NAVY

out = '/opt/projects/wh/docs/Lancer-Purge-Battle-Plan.docx'
doc.save(out)
print('saved', out)
