#!/usr/bin/env python3
"""Build Imperial Knights battle-plan Word docs (one per list) for the meta-tested
slate. Pulls the live roster/points (incl. validated allies) + Purge layouts.

    PYTHONPATH=src python3 tools/gen_knights_docx.py
-> docs/Knights-C-Battle-Plan.docx, docs/Knights-A-Battle-Plan.docx
"""
import os
os.environ.setdefault("WH_FACTION", "knights")
import yaml
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from wh import army as army_mod, data

STEEL = RGBColor(0x27, 0x3A, 0x5B)   # Knight house blue
GOLD = RGBColor(0x8A, 0x6D, 0x1F)

# shared per-archetype plans (both lists lock Purge the Foe)
PLANS = [
    ('Salamanders flamer-brick', 'Consecrate (vs Recon) / Meatgrinder (vs Purge)',
     'Delete the 2 Land Raiders T1-2 (Volcano/harpoon/thermal/spears), then Aggressors/Bladeguard with Avenger + sweeps. Knights shrug pyreblasters (T9-13 vs S5 AP-1).'),
    ('Ork Kult of Speed (dakka)', 'Punishment',
     'Anti-tank the Kill Rigs + Wazdakka; Avenger + sweeps for Deffkoptas/Flash Gitz (volume, not Volcano waste). Grab their home (8 VP) — fast but fragile.'),
    ('Ork Green Tide (100 Boyz)  ★ the test', 'Unstoppable Force',
     'The anti-horde game. Flamers/Avenger/sweeps clear Boyz; save harpoon/Volcano for Ghazghkull. HOLD with OC10 Knights (Boyz can’t shift a 400-pt Knight), SCREEN the charge with Armigers. You won’t table 100 Boyz — win on kills + holding while grinding the tide.'),
    ('Necron Monoliths (T13, OC8, slow)', 'Punishment',
     'Knights KILL Monoliths (Volcano S18 wounds on 3+, harpoon S24 deletes one) — the thing Sisters couldn’t. Focus-fire one down/turn; grab their home (8 VP) + objectives the M8" pyramids can’t reach.'),
    ("Necron C'tan (4× T11, −1 Damage)", 'Meatgrinder (vs Purge) / Consecrate (vs Recon)',
     'Can’t table 4 C’tan (4++ + −1 Damage). Use VOLUME (Avenger), not Volcano (AP-5 wasted on the invuln). Kill the killable enablers (Lokhust, Lychguard, Reanimator, characters); farm kill-differential, grab their home, feed the C’tan nothing.'),
    ('Dark Angels Deathwing (2+/4++, OC1)', 'Unstoppable Force',
     'Bricks are W4 — Avenger + sweeps are efficient; Volcano/harpoon is overkill waste. Chip an exposed brick, but WIN by out-holding OC1 bricks: your OC10 Knights + Armigers take more objectives than three OC1 bricks. Screen the Deep Strike.'),
]

LISTS = {
    'C': dict(
        file='examples/knights-C-horde-hardened.yaml', out='docs/Knights-C-Battle-Plan.docx',
        title='HORDE-HARDENED', at=27, horde=32,
        identity='The meta-robust pick. Strongest against the two matchups that actually beat '
                 'Knights — the Green Tide and C’tan-volume — while still deleting any single '
                 'tough target with the S24 Thundercoil harpoon. Valiant (DOMINUS) + Warden + Crusader '
                 'give three anti-horde big guns; 3 Warglaives screen and sweep.',
        notes={'Ork Green Tide (100 Boyz)  ★ the test': 'THIS LIST’S matchup — ~35 dead Boyz/turn + 4 Armiger screens. Field C here.',
               "Necron C'tan (4× T11, −1 Damage)": 'C’s Avenger volume out-damages Volcano into the 4++.',
               'Ork Kult of Speed (dakka)': 'C is fine but lacks A’s Lancer to run down the fast units.'}),
    'A': dict(
        file='examples/knights-A-lancer-aggressive.yaml', out='docs/Knights-A-Battle-Plan.docx',
        title='LANCER AGGRESSIVE', at=36, horde=23,
        identity='The aggressive, mobile pick. The Cerastus Lancer (4++ full invuln, M14) is the '
                 'most durable single model vs the AP-4 meta AND the fastest threat; the Castellan’s '
                 'Volcano gives max single-target punch. Balanced anti-horde, slightly softer into the '
                 'pure Green Tide than C.',
        notes={'Ork Kult of Speed (dakka)': 'THIS LIST’S edge — the Lancer (M14) catches and deletes the fast units.',
               'Necron Monoliths (T13, OC8, slow)': 'A’s Volcano + Lancer = the most single-target punch to crack pyramids.',
               'Ork Green Tide (100 Boyz)  ★ the test': 'A copes (~23 Boyz/turn) but C is the specialist — lean on OC10 holding + screens.'}),
}


def build_doc(cfg):
    doc = Document()
    n = doc.styles['Normal']; n.font.name = 'Calibri'; n.font.size = Pt(10.5)

    def h(text, level=1):
        p = doc.add_heading(text, level=level)
        for r in p.runs: r.font.color.rgb = STEEL if level == 1 else GOLD
        return p

    def para(*segs, after=4, italic=False):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
        for s in segs:
            r = p.add_run(s[0] if isinstance(s, tuple) else s)
            r.bold = isinstance(s, tuple) and s[1]; r.italic = italic
        return p

    def bullet(*segs):
        p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(2)
        for s in segs:
            r = p.add_run(s[0] if isinstance(s, tuple) else s); r.bold = isinstance(s, tuple) and s[1]
        return p

    t = doc.add_heading('IMPERIAL KNIGHTS', 0); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in t.runs: r.font.color.rgb = STEEL
    sub = para((f'List {list(LISTS.keys())[list(LISTS.values()).index(cfg)]} · {cfg["title"]}  ·  '
                f'LOCK: Purge the Foe', True)); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(cfg['identity'], italic=True, after=10)

    # 1. list
    h('1 · The Army List', 1)
    a = army_mod.build_file(cfg['file'])
    para((f'{a.name}  —  {a.points}/{a.points_limit} pts  (Knight portion + validated Agents ally)', True))
    para(('Detachments: Valourstrike Lance (2DP) + Dominus Foebreakers (1DP). Disposition: Purge the Foe.', False), after=6)
    for l in a.lines:
        label = l.datasheet + (' (ally)' if l.ally else '')
        extra = ('  ·  ' + ', '.join(l.wargear) if l.wargear else '') + (f'  ·  {l.enhancement}' if l.enhancement else '')
        bullet((f'{"%dx " % l.count if l.count > 1 else ""}{label}', True), (f'  {l.total} pts{extra}', False))
    para((f'Role balance: anti-tank {cfg["at"]} (vs a 16W Land Raider — redundant, all Knight lists overkill it) · '
          f'anti-horde {cfg["horde"]} Boyz/turn.', False), italic=True, after=8)

    # 2. how knights fight
    h('2 · How Knights fight this meta', 1)
    bullet(('Anti-tank is redundant; the axes are ANTI-HORDE + DURABILITY.', True),
           ' Every Knight gun overkills the meta’s vehicles (S18/S24 vs 16-22W); Monoliths die too. '
           'Games are lost to hordes you can’t clear and invuln-spam you can’t table — not to tanks.')
    bullet(('Hold with OC10 WHILE shooting.', True), ' A 400-pt Knight sits on an objective and fires 72" — '
           'holding IS the plan, not a compromise. Centre terrain blocks turn-1 LOS → advance to lanes.')
    bullet(('Durability vs the AP-4 meta:', True), ' Blessed Plate (T13) drops death-ray/lascannon damage a third; '
           'the Lancer’s 4++ (if fielded) caps AP-4 in melee AND at range. Armigers (T9) are the fragile link — '
           'screen with them, don’t expose them to focused melta.')
    bullet(('Judge big guns by target WOUNDS/model.', True), ' Volcano/harpoon into a W1 Boy or W4 Terminator wastes '
           '~7-9 damage — use Avenger/sweeps on multi-model units, save the big AP for tough single targets.')
    bullet(('Navigator locks your home.', True), ' Hidden + anti-Deep-Strike bubble = near-unremovable; frees an Armiger to push.')

    # 3. mission by opponent
    h('3 · Mission — set by the OPPONENT’s disposition (you’re locked Purge)', 1)
    for opp, mis in [('Take-and-Hold (Green Tide, Deathwing)', 'Unstoppable Force (3/kill + 4/obj + 5/central)'),
                     ('Disruption (Monoliths, Ork Kult)', 'Punishment (5/condemned + hold + 8 for their home)'),
                     ('Purge the Foe (C’tan, Salamanders)', 'Meatgrinder (kill-differential + 5 for their home)'),
                     ('Reconnaissance (Salamanders, C’tan)', 'Consecrate (needs Objective ACTIONS — weak for Knights; hold hard instead)')]:
        bullet((opp + ' → ', True), mis)

    # 4. deployment
    h('4 · Deployment (Purge-the-Foe layouts)', 1)
    para(('Universal: 6 objectives (2 home / 2 central-NML / 2 expansion), 16 obscuring terrain areas, '
          'centre packed (no turn-1 cross-board LOS). Big Knights hold centre + mid with LOS; Navigator on '
          'home; Armigers screen the big Knights and contest the flanks/expansion.', False), after=6)
    lay = yaml.safe_load(open('data/layouts/purge-the-foe.yaml'))
    for m in lay['matchups']:
        para((f"vs {m['vs'].replace('-', ' ').title()}  (you: {m.get('your_mission','')})", True), after=2)
        for L in ('A', 'B', 'C'):
            d = m['layouts'][L]
            bullet((f'Layout {L}: ', True), f"{d['deployment']} — {d.get('note','')}")

    # 5. per-archetype
    h('5 · Per-archetype plans', 1)
    for name, mission, targets in PLANS:
        h(name, 2)
        para(('Mission: ', True), mission, after=2)
        bullet(('Plan: ', True), targets)
        if name in cfg['notes']:
            bullet(('This list: ', True), cfg['notes'][name])

    doc.save(cfg['out'])
    print(f"wrote {cfg['out']}  ({a.points}/{a.points_limit} pts, {len(a.lines)} units)")


for cfg in LISTS.values():
    build_doc(cfg)
