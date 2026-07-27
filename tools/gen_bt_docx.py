#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Black Templars 'Send Help' FIXED Runbook (Word) from tools/bt_data.py + mc_bt_sim.py.
   PYTHONPATH=tools:src python3 tools/gen_bt_docx.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import bt_data as D
import doc_versions as V

DOCKEY = "bt-runbook"
OUT = V.out_path(DOCKEY)
STEEL = RGBColor(0x20, 0x20, 0x28)   # BT black
SILVER = RGBColor(0x6B, 0x70, 0x78)  # BT steel
VERD = {"FAV": RGBColor(0x2E, 0x7D, 0x32), "EVEN": RGBColor(0x2E, 0x7D, 0x32), "COIN": RGBColor(0xB8, 0x86, 0x00),
        "UNFA": RGBColor(0xC0, 0x39, 0x00), "HARD": RGBColor(0xB0, 0x20, 0x00)}


def vcol(v):
    for k in VERD:
        if v.upper().startswith(k):
            return VERD[k]
    return STEEL


def h(doc, text, size, color=STEEL, before=8, after=4):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    return p


def kv(doc, k, v):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    p.add_run(k + ": ").bold = True; p.add_run(v)


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def main():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10)
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("BLACK TEMPLARS — 'SEND HELP' FIXED RUNBOOK"); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = STEEL
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run(f"{D.EVENT}\nGenerated {D.GENERATED}").italic = True
    vp = doc.add_paragraph(); vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = vp.add_run(V.cover_line(DOCKEY)); vr.bold = True; vr.font.size = Pt(11); vr.font.color.rgb = SILVER

    h(doc, "THE FIXED LIST", 15)
    kv(doc, "Detachments", D.DETACHMENTS)
    kv(doc, "Disposition", D.DISPOSITION)
    kv(doc, "Points", D.LIST_TOTAL)
    doc.add_paragraph(D.IDENTITY).italic = True
    for name, cnt, role, pts in D.UNITS:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{cnt}x {name} ({pts}) — ").bold = True; p.add_run(role)

    h(doc, "THE TAPESTRY — how the fixed list wins", 15)
    for name, desc in D.RULES:
        p = doc.add_paragraph(style="List Bullet"); p.add_run(name + " — ").bold = True; p.add_run(desc)
    h(doc, "Mindset", 12, color=RGBColor(0, 0, 0), before=6)
    doc.add_paragraph(D.MINDSET)
    h(doc, "Realistic record", 12, color=RGBColor(0, 0, 0), before=6)
    doc.add_paragraph(D.RECORD_NOTE)
    for band, lists in D.BANDS.items():
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        p.add_run(band + ": ").bold = True; p.add_run(", ".join(lists))

    # sim verification
    doc.add_page_break()
    h(doc, "SIM VERIFICATION (data-driven)", 15)
    doc.add_paragraph("Games/archetype via tools/mc_bt_sim.py. Spike = the delivered Sword Brethren spike "
                      "(crit-on-5 Lethal Hits, S7 vs Monster/Vehicle) onto the priority target; Shoot = BT's thin "
                      "Oath-boosted shooting. Delivery-dependent (2 buses). Indicative — coin-flips are the "
                      "matchups to practice.").italic = True
    try:
        import mc_bt_sim
        rows = mc_bt_sim.results(2000)
    except Exception as e:
        rows = []
        doc.add_paragraph(f"(sim unavailable: {e})")
    if rows:
        tbl = doc.add_table(rows=1, cols=6); tbl.style = "Light Grid Accent 1"; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, x in enumerate(("Archetype", "Prev.", "Verdict", "Spike", "Shoot", "BT win%")):
            tbl.rows[0].cells[i].paragraphs[0].add_run(x).bold = True
        for x in rows:
            c = tbl.add_row().cells
            c[0].text = x["archetype"]; c[1].text = str(x["prev"]); c[2].text = x["verdict"]
            c[3].text = str(x["spike"]); c[4].text = str(x["shoot"])
            rr = c[5].paragraphs[0].add_run(f"{x['win']}%"); rr.bold = True
        for col, wd in zip(tbl.columns, (2.6, 0.7, 1.0, 0.9, 0.9, 0.9)):
            for cell in col.cells:
                cell.width = Inches(wd)
        prev = sum(x["prev"] for x in rows); wr = sum(x["prev"] * x["win"] for x in rows) / prev
        doc.add_paragraph(f"Prevalence-weighted BT win rate across the top meta ≈ {wr:.0f}% — a SWINGY MID/LOW-"
                          f"tier melee-spike army (below a strong build's ~57%). Deletes what a spike REACHES "
                          f"(vehicles/monsters/elites); loses to castled shooting (T'au), true hordes, and pure "
                          f"speed. Delivery-dependent.").italic = True

    # version history (incremental-improvement tracker)
    doc.add_page_break()
    h(doc, "VERSION HISTORY — incremental improvements", 15)
    doc.add_paragraph("Same Sword-Brethren-spike engine, iterated on DELIVERY. Win% per archetype for each list "
                      "version (tools/mc_bt_sim.py, all configs). The documented list is the last column.").italic = True
    try:
        import mc_bt_sim
        names, hrows, weighted = mc_bt_sim.history(2000)
    except Exception as e:
        names, hrows, weighted = [], [], {}
        doc.add_paragraph(f"(history unavailable: {e})")
    if names:
        tbl = doc.add_table(rows=1, cols=2 + len(names)); tbl.style = "Light Grid Accent 1"; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        heads = ["Archetype", "Prev."] + names
        for i, x in enumerate(heads):
            tbl.rows[0].cells[i].paragraphs[0].add_run(x).bold = True
        for r in hrows:
            c = tbl.add_row().cells
            c[0].text = r["archetype"]; c[1].text = str(r["prev"])
            for j, n in enumerate(names):
                c[2 + j].text = f"{r[n]}%"
        wr = tbl.add_row().cells
        wr[0].paragraphs[0].add_run("PREVALENCE-WEIGHTED").bold = True
        for j, n in enumerate(names):
            wr[2 + j].paragraphs[0].add_run(f"{weighted[n]}%").bold = True
    for name, desc in D.CHANGELOG:
        p = doc.add_paragraph(style="List Bullet"); p.add_run(name + " — ").bold = True; p.add_run(desc)

    # battle plans
    doc.add_page_break()
    h(doc, "BATTLE PLANS", 16)
    for m in D.MATCHUPS:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.keep_with_next = True
        r = p.add_run(f"{m['faction']} — {m['archetype']}"); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = STEEL
        vp = doc.add_paragraph(); vp.paragraph_format.space_after = Pt(2)
        vr = vp.add_run(m["verdict"]); vr.bold = True; vr.font.color.rgb = vcol(m["verdict"])
        vp.add_run(f"   ·   Prevalence: {m['prev']}").italic = True
        kv(doc, "Deciding factor", m["deciding"])
        gp = doc.add_paragraph(); gp.paragraph_format.space_after = Pt(2); gp.add_run("Game plan:").bold = True
        bullets(doc, m["plan"])
        kv(doc, "Watch for", m["watch"])

    doc.add_page_break()
    h(doc, "APPENDIX — VERIFIED PROFILES (the numbers behind the sim)", 13)
    for wn, prof, note in D.VERIFIED_PROFILES:
        if wn.strip().startswith("---"):
            h(doc, wn.strip("- "), 12, color=RGBColor(0, 0, 0), before=8); continue
        p = doc.add_paragraph(style="List Bullet"); p.add_run(wn + ": ").bold = True
        p.add_run(prof + (f"  — {note}" if note else ""))

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    V.stamp_docx(doc, DOCKEY)
    doc.save(OUT)
    print("wrote", OUT, V.cover_line(DOCKEY), "-", len(D.MATCHUPS), "battle plans")


if __name__ == "__main__":
    main()
