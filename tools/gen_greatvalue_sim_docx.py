#!/usr/bin/env python3
"""Full-game simulation: my Knights (Freeblade Company, Priority Assets) vs the friend's
"Great Value" Imperial Fists (Emperor's Shield + Librarius Conclave, Purge the Foe).

Model-by-model, phase-by-phase, turn-by-turn to a conclusion. Demonstrates a complete
game of 11E 40k -- not just dice, but sequencing, missions, secondaries, CP, reserves,
Oath, disciplines, stratagems, and the objective/kill scoring race.

    PYTHONPATH=src python3 tools/gen_greatvalue_sim_docx.py
 -> docs/Great-Value-vs-Knights-Full-Game-Simulation.docx
"""
import os
os.environ.setdefault("WH_FACTION", "knights")
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

STEEL = RGBColor(0x27, 0x3A, 0x5B)   # Knight blue
GOLD  = RGBColor(0x8A, 0x6D, 0x1F)
FIST  = RGBColor(0x8B, 0x1A, 0x1A)   # Imperial Fists red
GREY  = RGBColor(0x55, 0x55, 0x55)

doc = Document()
base = doc.styles["Normal"]
base.font.name = "Calibri"; base.font.size = Pt(10.5)

def H(txt, lvl=1, color=STEEL):
    h = doc.add_heading(txt, level=lvl)
    for r in h.runs: r.font.color.rgb = color
    return h

def P(txt="", bold=False, italic=False, color=None, size=None, align=None, space=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    if align: p.alignment = align
    r = p.add_run(txt); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    return p

def rich(parts, space=4):
    """parts = list of (text, {bold,italic,color,size})."""
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space)
    for txt, fmt in parts:
        r = p.add_run(txt)
        r.bold = fmt.get("bold", False); r.italic = fmt.get("italic", False)
        if fmt.get("color"): r.font.color.rgb = fmt["color"]
        if fmt.get("size"): r.font.size = Pt(fmt["size"])
    return p

def bullet(txt, lvl=0, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet" if lvl == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
        p.add_run(txt)
    else:
        p.add_run(txt)
    return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v)); r.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t

def hr():
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    r = p.add_run("─" * 60); r.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

# ============================================================ TITLE
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("A COMPLETE GAME OF WARHAMMER 40,000 (11E)"); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = STEEL
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Imperial Knights (Freeblade Company) vs. “Great Value” Imperial Fists")
r.font.size = Pt(13); r.italic = True; r.font.color.rgb = GREY
s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s2.add_run("Lone Star Open practice — full model-by-model, phase-by-phase simulation to conclusion")
r.font.size = Pt(10); r.font.color.rgb = GREY
P()
P("How to read this document", bold=True, color=GOLD)
P("This is a simulation of an entire game, played competently from BOTH sides. Every battle round "
  "has my turn and his turn; every turn walks the phase sequence (Command → Movement → Shooting "
  "→ Charge → Fight → end-of-turn scoring). Dice results are stated as expected values (EV) from "
  "the mathhammer engine, rounded to plausible whole-model outcomes — i.e. what you should PLAN around, "
  "with variance called out where it swings the game. A running VP tally closes each round. The point is "
  "to show the game is won by SEQUENCING and the MISSION, not by dice.", size=10)

# ============================================================ 1. THE TWO ARMIES
H("1. The two armies — and their tapestries", 1)

P("His list: “Great Value” — Imperial Fists, 1985 pts", bold=True, color=FIST)
P("Detachments (3 DP): Emperor’s Shield (2 DP) + Librarius Conclave (1 DP). Disposition: Purge the Foe.", italic=True, size=9)
table(["Unit (attached)", "Key profile", "The threads that matter"],
[["Darnath Lysander + Ancient in Term. + 10 Assault Terminators (TH/SS)",
  "T5 W4 2+/4++ ×11, M5, OC2. TH: S8 AP-2 D2 A5 Dev; Fist of Dorn S10 AP-3 D3 A5 Dev",
  "Icon of Obstinacy (−1 to wound the unit when S≥T5), Rampart (2+ inv once), Inspiring Commander (non-Character Terminators are OC2 WHILE NOT BATTLE-SHOCKED → brick OC22; base OC12, and a Battle-shock drops it back to OC12), Champion of the Feast (+1 A), Terminatus Assault (Battle-shock on engage). Wrath of Dorn: FULL wound re-roll vs Oath target."],
 ["Bladeguard Ancient + 6 Bladeguard Vets",
  "T4 W3 3+/4++, MC power weapon S5 AP-2 D2",
  "Malodraxian Standard (−1 to wound when S>T4), Bladeguard (pick Swords=rr hit-1 or Shields=rr inv-1 each Fight). Very hard to shift."],
 ["Librarian [Temporal Corridor] + 10 Sternguard",
  "T4 W2 3+, bolt rifle S4 AP-1 D1 A2 Assault/Dev/Heavy/RF1",
  "PSYKER unit (gets the round’s Discipline). Sternguard Focus (FULL wound re-roll vs Oath). Temporal Corridor: Deep Strike if Telekinesis; bail to Reserves at end of my Fight phase if unengaged. THE unsaveable Dev-Wound chip — but SOFT (T4 W2, no invuln)."],
 ["Librarian in Term. [Fusillade] + 10 Terminators (2 cyclone)",
  "T5 W3 2+/4++, cyclone krak S9 AP-2 D6 A2; storm bolter S4 RF2; Teleport Homer",
  "PSYKER unit. Fusillade (Lethal Hits; +Sustained under Pyromancy). Fury of the First (+1 Hit vs Oath, datasheet). OC22 un-shocked (Inspiring Commander), OC12 base. The real ranged Knight-threat — but slow (M5)."],
 ["2× Vanguard Vets w/ Jump Packs (5 ea)",
  "T4 W2 3+/4++, M12, jump",
  "Vanguard Assault (Lethal Hits on the charge). Fast harassers / action-monkeys."],
 ["Intercessor Squad (5)", "T4 W2 3+, OC2", "Objective Secured (sticky home). Target Elimination (+2 A once)."],
 ["2× Land Speeder", "T8 W9 3+, M14 FLY, OC3", "Multi-melta S9 AP-4 D6 A2 Melta 2 + stormfury S12 AP-3 D6+1. Purgation Run (shoot then move D6”). NOTE: no innate Deep Strike — arrives from Strategic Reserves (board edge)."]],
widths=[2.0, 2.2, 3.0])
P("His engine: OATH OF MOMENT.", bold=True, color=FIST)
P("Every Command phase he names one of my units; his whole army re-rolls Hit rolls AND gets +1 to Wound "
  "against it (the +1 because he’s mono-Codex Imperial Fists). Layered with Wrath of Dorn (re-roll Wound-of-1 "
  "army-wide, full for Lysander’s unit), the unit synergies above, and the round’s Librarius Discipline "
  "(Divination re-roll 1s / Pyromancy +1 AP+Sustained / Telekinesis Deep-Strike-enable & −1 S defence), his "
  "fully-converged fire onto one Oathed Knight is ~34 damage — a dead Castellan. The counter is written into "
  "the numbers: ~10 of that 34 is the SOFT Sternguard. Kill a leg early and the alpha collapses.", size=10)

P()
P("My list: Knights v2 — Freeblade Company, 1985 pts", bold=True, color=STEEL)
P("Detachment (3 DP): Freeblade Company. Disposition: Priority Assets. Army rule: Code Chivalric (Knight vows).", italic=True, size=9)
table(["Unit", "Key profile", "Role & threads"],
[["Knight Crusader [Hunter’s Eye]",
  "T11 W26 3+/5++(ranged), OC10. Avenger gatling S6 AP-2 D2 A18; thermal cannon S12 AP-4 Melta6; RFBC S10",
  "The anti-SUPPORT gun. Avenger deletes Sternguard; Hunter’s Eye = Ignores Cover (nothing hides). Thermal/RFBC chip Terminators."],
 ["Knight Castellan [Sanctuary]",
  "T12 W28 3+/5++→4++full(Sanctuary), OC10. Volcano lance S18 AP-5 D6+8; plasma decimator; shieldbreaker; siegebreaker",
  "Durable anchor + long-range sniper. Volcano one-shots Land Speeders and chunks Lysander/characters. Sanctuary = full 5++ (melee too). The likely Oath target — built to survive it."],
 ["Cerastus Knight Lancer",
  "T11 W28 3+/4++ FULL, M14, OC10. Shock lance S20 AP-3 D8 (Shock Charge: Tank Shock 0CP)",
  "The flanker. Hunts Land Speeders, dives his backline, one-shots characters, and grabs his home late. 4++ full = tanks a focus if needed."],
 ["2× Armiger Warglaive",
  "T9 W14 3+/5++, M12, OC6. Thermal spear S12 AP-4 Melta4",
  "Forward screens / pins / BAIT (soak the Oath so it isn’t a big Knight). Anti-Speeder melta. Bondsman recipients. Expendable at 140 ea."],
 ["Armiger Helverin",
  "T9 W14 3+/5++, M12, OC6. 2× Armiger autocannon S9 AP-1 D3 A4",
  "Backfield fire — anti-Sternguard / anti-Speeder / anti-chaff volume."],
 ["Navigator (ally)", "T3 W4, character", "12” anti-Deep-Strike dome — denies Temporal Corridor / Dropship Extraction / Land Speeder reserve re-drops near my scoring."],
 ["Battle Sisters Squad (ally, split)", "T3 W1 3+, OC2/model", "Requisitioned; splits → OC10 foot half locks my home under the dome; melta half rides the Immolator."],
 ["Immolator (ally)", "T9 W11 3+, twin multi-melta", "Twin MM (NOT flamers vs 2+/4++). Purge & Cleanse: strips Benefit of Cover on a target for my WHOLE army’s shooting."]],
widths=[1.7, 2.5, 3.0])
P("My engine: FEEL NO PAIN + the MISSION.", bold=True, color=STEEL)
P("Freeblade Company’s Knights of Legend gives every Knight Feel No Pain 6+ and regains 1 lost wound per "
  "Command phase. FNP is the ONLY thing that touches his unsaveable Sternguard Dev Wounds, and the regen "
  "un-does his chip. I do not try to out-kill unkillable bricks — I picked Priority Assets to win the "
  "OBJECTIVE race with speed + spread, and to steal his home for 10 VP at the end.", size=10)

# ============================================================ 2. MISSION / DEPLOYMENT
H("2. Mission, deployment & plan", 1)
rich([("Primary missions (asymmetric): ", {"bold":True}),
      ("my Priority Assets vs his Purge the Foe → I play ", {}),
      ("Vital Link", {"bold":True,"color":STEEL}),
      (", he plays ", {}),
      ("Destroyer’s Wrath", {"bold":True,"color":FIST}), (".", {})])
table(["", "Vital Link (me)", "Destroyer’s Wrath (him)"],
[["End of MY/HIS turn", "2 VP control a central objective; +1 per operation marker there", "3 VP if I destroyed 1+ enemy unit this turn"],
 ["Round 2+, end of Command", "4 VP control a non-home objective; +4 if one is central", "4 VP control a non-home objective; 6 VP control MORE objectives than opponent"],
 ["Round 2+, end of turn", "—", "4 VP destroyed more enemy units this turn than he lost last turn"],
 ["End of battle", "10 VP control opponent’s home", "—"]],
widths=[1.7, 2.6, 2.6])
P("Reading the clash: he scores by KILLING (easy — he’ll pop a screen most turns) and by OUT-CONTROLLING "
  "objectives (the 6-VP swing). I score by HOLDING objectives and by taking his home late. Both primaries cap "
  "at 15 VP/round. So my job is: (a) deny him the objective-control clauses so he’s stuck at ~7/round from "
  "kills, (b) hold enough to bank ~10–13/round myself, (c) steal his home for the closing 10.", size=10)

P("Deployment — Crucible of Battle (diagonal). 5 objectives:", bold=True)
bullet("A = HIS home (his corner). B = MY home (my corner). C = CENTRE. D = left no-man’s-land. E = right no-man’s-land.")
P("I win the roll-off and take the first turn (seize objectives; delete his exposed Sternguard before he can "
  "hide them). He deploys reactively and holds mobile pieces back.", size=10)
rich([("My deployment: ", {"bold":True, "color":STEEL}),
      ("Castellan centre-left (LOS to C + no-man’s), Crusader centre-right (LOS to E/mid). Battle Sisters SPLIT: "
       "OC10 foot half on my home B under the Navigator’s 12” dome; melta half embarked in the Immolator, "
       "centre. 2 Warglaives forward as screens toward C/D. Helverin backfield near B with LOS to mid. Lancer far "
       "right, poised to run up the flank at E and his home. Navigator by B.", {})])
rich([("His deployment: ", {"bold":True, "color":FIST}),
      ("TH/SS brick (Lysander) centre-front behind a ruin. Bladeguard near C. Sternguard mid with firing lines "
       "(they must see to shoot — exposed). One Vanguard squad left flank, one right. Intercessors on his home A. "
       "IN RESERVE: the cyclone Terminator Squad (Deep Strike, to a Teleport Homer) and both Land Speeders "
       "(Strategic Reserve). He places 2 Teleport Homer tokens mid-board near C and E — OUTSIDE my dome.", {})])
P("CP: each player starts 0 and gains 1 at the start of each of their Command phases. Secondaries: both play "
  "Tactical (draw 2 each turn). His danger picks: Bring It Down (my Knights are all VEHICLE/TITANIC) and "
  "Assassination (my Knights are CHARACTERS). Mine: Behind Enemy Lines, Engage on All Fronts, Storm Hostile "
  "Objective, Area Denial — all rewarding my speed.", size=9, color=GREY)

# ---------- helper for a running tally line
def tally(me, him, note=""):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
    r = p.add_run("   RUNNING TOTAL — "); r.bold = True
    r = p.add_run(f"Knights {me}"); r.bold = True; r.font.color.rgb = STEEL
    r = p.add_run("  vs  ");
    r = p.add_run(f"Imperial Fists {him}"); r.bold = True; r.font.color.rgb = FIST
    if note:
        r = p.add_run("   " + note); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY

def turnhead(txt, color):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(txt); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = color

def phase(name, body):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(name + "  "); r.bold = True; r.font.color.rgb = GOLD; r.font.size = Pt(10)
    r = p.add_run(body); r.font.size = Pt(10)

# ============================================================ ROUND 1
H("Battle Round 1", 1)
turnhead("▶ My Turn 1 — Knights (seize + decapitate his ranged)", STEEL)
phase("Command.", "Gain 1 CP (1 total). No Freeblade regen needed (undamaged). Code Chivalric vow: “We pledge "
  "to reap a great tally” is unwise vs unkillable bricks — I take the objective-holding vow instead. Priority "
  "Assets scores at end of turn.")
phase("Movement.", "Both Warglaives advance ~9–10” to screen C and D (bodies between his brick and my big "
  "Knights). Immolator (melta Sisters aboard) moves 12” to sit ON centre C. Lancer advances up the right (M14 + "
  "advance = ~20”) toward E, staying just out of his charge threat. Castellan & Crusader shuffle for clean LOS to "
  "the Sternguard, holding near centre. Sisters foot half + Navigator stay on B.")
phase("Shooting.", "DECAPITATE THE RANGED. Crusader [Hunter’s Eye, Ignores Cover] into the 10 Sternguard: "
  "Avenger gatling 18 shots → 12 hits → S6 vs T4 wounds on 3+ → 8 wounds → AP-2 vs 3+ = 5+ save → "
  "~5 fail → D2 → ~10 damage = 5 Sternguard down; thermal cannon + RFBC add ~4 more → 9 of 10 dead. "
  "Helverin’s 2 autocannons (8 shots) finish the last Sternguard. STERNGUARD WIPED T1. Castellan plasma decimator "
  "(supercharge, 6 Blast shots) + siegebreaker into the Bladeguard → −1-to-wound (Malodraxian) + 4++ → only "
  "2 Bladeguard die; Volcano held (no good target worth S18 yet). Immolator holds fire / stays on C.")
phase("Charge / Fight.", "None — I do not charge the brick. My screens sit.")
phase("End of my turn — Vital Link.", "I control centre C (Immolator OC3 + a Warglaive, uncontested) → 2 VP. "
  "(No Maintain Control action this turn — I’d rather shoot.) Secondary: Engage on All Fronts (units in 3 table "
  "quarters) → 3 VP.")
tally("5", "0", "Primary 2 + Secondary 3. His single best Knight-killer (Sternguard, ~10 unsaveable/turn) is GONE on turn one.")

turnhead("▶ His Turn 1 — Imperial Fists (advance the wall)", FIST)
phase("Command.", "Gain 1 CP. Oath of Moment: names my Castellan (habit) — but he has almost nothing on the "
  "board that shoots it now. Librarius Discipline: with his psyker units dead/in-reserve, he picks Telekinesis to "
  "prime a Deep-Strike round next turn.")
phase("Movement.", "TH/SS brick advances (M5 + advance ~9”) toward centre C. Bladeguard advance to C’s edge. "
  "Both Vanguard jump 12” — the right-flank squad lands near my forward Warglaive. Intercessors sit on home A.")
phase("Shooting.", "Thin: Intercessor bolt rifles + Vanguard/ Bladeguard pistols into a Warglaive → T9, ~3 wounds "
  "(FNP 6+ shrugs 1) → Warglaive at ~11W. Nothing else reaches.")
phase("Charge.", "Right Vanguard (5) declare a charge on the forward Warglaive; make it. Vanguard Assault → "
  "Lethal Hits this turn.")
phase("Fight.", "Vanguard: ~15 attacks, Lethal Hits, S4-5 AP-1/-2 into T9 → ~4 wounds (FNP shrugs ~1) → "
  "Warglaive to ~7W, survives. Warglaive swings the reaper chain-cleaver back → 2 Vanguard dead. The Warglaive is "
  "now pinned by 3 Vanguard — fine; it can fall back next turn and still shoot (Super-heavy Walker).")
phase("End of his turn — Destroyer’s Wrath.", "He destroyed 0 of my units this turn → 0 VP. Secondary: "
  "Area Denial / Cleanse → 3 VP.")
tally("5", "3", "His wall is intact but SLOW and his ranged is gone. He’s betting everything on turns 2–4.")

# ============================================================ ROUND 2
H("Battle Round 2 — the alpha strike lands (and fizzles)", 1)
turnhead("▶ My Turn 2 — Knights (bank the board, kite the brick)", STEEL)
phase("Command.", "Gain 1 CP (2 total). Freeblade regen: the pinned Warglaive +1 W (7→8). Vital Link round-2+ "
  "clause fires at end of Command — but I want to move first to secure objectives, so it will resolve on my "
  "positions this turn. I hold C (Immolator) and D (Warglaive) → 4 VP (non-home) + 4 VP (C is central) = 8 VP.")
phase("Movement.", "KITE. Castellan & Crusader reposition AWAY from the advancing brick while keeping LOS — the "
  "brick is M5, I am faster; it will never catch a Knight that respects it. The pinned Warglaive FALLS BACK from the "
  "Vanguard (Super-heavy Walker lets it still shoot). Lancer surges up the right to threaten E and his backfield. "
  "Immolator holds C; Sisters melta half stays aboard.")
phase("Shooting.", "His cyclone Terminators + Speeders are still in reserve, so I shoot what’s here. Crusader + "
  "Helverin + the fallen-back Warglaive into the Bladeguard: through −1-to-wound + 4++ they’re stubborn — "
  "~3 more Bladeguard die (unit down to 1 model + Ancient). Castellan Volcano into the TH/SS brick just to bank chip "
  "→ Icon of Obstinacy (−1 wound) + 4++ → ~1 Terminator + wounds. I deliberately do NOT dump melta into the "
  "brick’s 4++ — that’s the trap.")
phase("Charge / Fight.", "None. Discipline over aggression.")
phase("End of my turn — Vital Link.", "End-of-turn control of central C → 2 VP. Maintain Control action was "
  "skipped again (shooting mattered more). Round total = 8 (Command) + 2 (end) = 10 VP, under the 15 cap. Secondary: "
  "Behind Enemy Lines (Lancer approaching his half) pending; Engage → 4 VP.")
tally("19", "3", "Primary 10 + Secondary 4. I’m banking the objective race while his army is still walking.")

turnhead("▶ His Turn 2 — Imperial Fists (THE alpha)", FIST)
phase("Command.", "Gain 1 CP (he’s been hoarding — ~3 CP). Oath: names my CASTELLAN (juiciest — also "
  "feeds Bring It Down + Assassination). Discipline: Divination (re-roll 1s to Hit & Wound) to sharpen the drop.")
phase("Movement / Reserves.", "Cyclone Terminator Squad DEEP STRIKES onto the Teleport Homer near E — it must be "
  ">9” from my models and can’t be inside the Navigator dome (which covers B and my centre-left), so it lands "
  "off my right, ~9” from the Crusader/Lancer. Both Land Speeders arrive from Strategic Reserve on his right board "
  "edge and move 14” up — but at >9”, so no Melta bonus this turn. The TH/SS brick keeps grinding forward.")
phase("Shooting — the converged alpha into the Oathed Castellan.", "Cyclone Terminators (Fury +1 Hit, Oath "
  "re-roll Hit + +1 Wound, Wrath re-roll Wound-1, Fusillade Lethal, Divination) ≈ 9 dmg. Land Speeders (2 MM + "
  "2 stormfury, Oath, >9” so no Melta) ≈ 8 dmg. RAW ≈ 17 — and here is the whole game: because I killed "
  "the STERNGUARD on turn 1, the ~10 unsaveable Dev Wounds simply aren’t in this total. Freeblade FNP 6+ shaves "
  "~3 more → ~14 damage. Castellan (28 W, Sanctuary) drops to ~14 W — bracketed, but ALIVE and still shooting.")
phase("Charge.", "The TH/SS brick reaches my line — but I screened, so it can only reach a WARGLAIVE, not a big "
  "Knight. It charges the Warglaive. Terminatus Assault → the Warglaive takes a Battle-shock test (2D6 vs Ld 7+): "
  "it fails (rolls 6) → OC0, no stratagems for it this turn. (I don’t care — it’s bait.)")
phase("Fight.", "The brick (not Oathed onto the Warglaive, but Champion-of-the-Feast attacks + hammers) ≈ 20+ "
  "damage into the T9 Warglaive → DEAD. But that is a 140-pt Armiger soaking a 450-pt brick’s entire activation, "
  "one board-corner away from anything that matters. The Warglaive’s Deadly Demise (D6) explodes on death → "
  "~2 mortal wounds spread onto the brick (1 Terminator chipped). Bait: executed.")
phase("End of his turn — Destroyer’s Wrath.", "Destroyed 1 unit (the Warglaive) → 3 VP. Round-2+ "
  "objective clauses: does he control a non-home objective? His brick is near C but my Immolator (OC3) + presence "
  "still contest — and his OC22 brick, IF it’s within range of C, would take it. Ruling it his: he controls C "
  "→ 4 VP; but he does NOT control MORE objectives than me (I hold D + E-adjacent + my home; he has C + his home) "
  "→ no 6 VP. ‘Destroyed more than I lost last turn’ (I lost 0 last turn) → 4 VP. Secondary: Bring It "
  "Down (killed a Warglaive = Vehicle) → 4 VP.")
tally("19", "18", "He’s caught up on kills — exactly what Destroyer’s Wrath rewards. But he spent his ALPHA and "
  "only bracketed a Castellan + ate a bait Armiger. His whole plan just under-performed.")

# ============================================================ ROUND 3
H("Battle Round 3 — the grind turns", 1)
turnhead("▶ My Turn 3 — Knights (kill the mobile legs, take the flank)", STEEL)
phase("Command.", "Gain 1 CP (3). Freeblade regen: Castellan 14→15 W; other chipped Knights +1. Vital Link "
  "round-2+ clause: I re-take/hold objectives at end of Command — if his brick sits on C, I pivot to hold D + E "
  "instead → 4 VP (non-home) but NOT the +4 central bonus this turn (he’s squatting C). I accept 4 and make it "
  "back on his home later.")
phase("Movement.", "The Lancer (M14) crashes into his right — lines up a charge on a Land Speeder. Crusader "
  "repositions to bear on the freshly-arrived cyclone Terminators. Castellan (bracketed but fine) kites further from "
  "the brick and lines up the Volcano on a Land Speeder. Helverin + surviving Warglaive hold D and screen. Immolator "
  "+ Sisters contest C’s edge.")
phase("Shooting.", "KILL THE MOBILE LEGS. Castellan Volcano into a Land Speeder: S18 AP-5 D6+8 → one-shot "
  "(~19 dmg vs 9 W) → Speeder #1 gone. Crusader (Ignores Cover) Avenger + thermal + RFBC into the cyclone "
  "Terminators (2+/4++ W3): ~6–7 wounds through the 4++ → 2 Terminators dead + the unit chipped. Helverin "
  "autocannons into Speeder #2 → ~5 dmg (to ~4 W).")
phase("Charge.", "Lancer declares on Land Speeder #2 — Shock Charge (Tank Shock, 0 CP): the Lancer’s Strength "
  "vs the Speeder inflicts mortal wounds on the charge → finishes Speeder #2. BOTH Land Speeders dead by end of "
  "T3. The Lancer is now deep on his right, one move from his home A.")
phase("Fight.", "Lancer’s shock lance (S20 AP-3 D8) is now free to threaten his backfield characters next turn.")
phase("End of my turn — Vital Link.", "Control D + E → already scored the 4 at Command; end-of-turn central "
  "control belongs to his brick this turn (0 central VP). I DO run a Maintain Control action on E with the Lancer’s "
  "wake / a Sister unit isn’t there… skip. Round primary = 4 VP. Secondary: Behind Enemy Lines (Lancer in his "
  "DZ) + Bring It Down (2 Speeders!) → 6 VP.")
tally("29", "18", "Primary 4 + Secondary 6. Both his mobile ranged legs are dead; his remaining threat is one chipped "
  "cyclone brick and the slow melee wall. The alpha is spent.")

turnhead("▶ His Turn 3 — Imperial Fists (squeeze the objectives)", FIST)
phase("Command.", "Gain 1 CP. Oath: names the Castellan again (finish it) OR the Lancer (it’s now deep in his "
  "lines threatening home). He Oaths the LANCER to protect his backfield. Discipline: Divination.")
phase("Movement.", "The TH/SS brick pivots to SQUAT the centre C (OC22 un-shocked — to out-control it I must pile ~OC23+ nearby, e.g. two OC10 Knights + an Armiger, OR Battle-shock it down to OC12). "
  "Bladeguard remnant + Ancient shuffle to contest D. The surviving Vanguard chase objectives. He uses Dropship "
  "Extraction? No — no upside now; he needs bodies ON objectives, not in reserve.")
phase("Shooting.", "Cyclone Terminators (now ~8 strong) into the Oathed Lancer: ~10 raw → FNP 6+ → ~8 → "
  "Lancer (28 W, 4++ full) to ~20 W — barely scratched. His melee wall has no ranged answer to a 4++ Knight in "
  "his backfield.")
phase("Charge.", "Bladeguard Ancient + last Bladeguard try to wall D; no charge into my Knights that profits him.")
phase("Fight.", "—")
phase("End of his turn — Destroyer’s Wrath.", "Destroyed 0 of my units this turn → 0 on that clause. "
  "Objectives: he controls C (brick OC22) + his home A + contests D → he may control MORE objectives than me this "
  "snapshot → 6 VP + 4 VP (non-home C) = 10 VP. ‘Destroyed more than I lost’ (I lost 0) → 0. Secondary: "
  "Cleanse → 3 VP. This is his big primary turn — the OC22 brick squatting the centre is real.")
tally("29", "31", "He surges on the objective-control clause — the OC22 brick on the centre is his one genuine "
  "board-control tool. I MUST break the ‘more objectives than me’ math next turn.")

# ============================================================ ROUND 4
H("Battle Round 4 — win the objective count, aim for his home", 1)
turnhead("▶ My Turn 4 — Knights (out-spread the wall)", STEEL)
phase("Command.", "Gain 1 CP (4). Regen tops up the Castellan/Lancer. Vital Link: I hold D + E + push my OC10 "
  "into range of MORE objectives than he can reach with two slow bricks. End-of-Command control: D + E (non-home) "
  "→ 4 VP (no central bonus — he still squats C).")
phase("Movement.", "SPREAD to win the count. Lancer sits ON his home A (OC10 — out-controls his 5 Intercessors, "
  "OC10>OC10? Intercessors sticky-held it, but my OC10 + killing/​contesting flips it if I clear them). Crusader + "
  "Helverin + Warglaive fan across D, E and centre-adjacent so that across the board I control B(home)+D+E and "
  "contest A — 3–4 objectives to his 1–2. Immolator + Sisters hold C’s edge to force his brick to choose.")
phase("Shooting.", "Clear his home for the Lancer: Crusader + Helverin delete the 5 Intercessors on A (soft T4 W2). "
  "Castellan Volcano/plasma into the cyclone Terminators → 2–3 more dead. I keep NOT feeding the brick.")
phase("Charge.", "Lancer charges the last Intercessor/​contests A cleanly → now I control HIS home objective A.")
phase("Fight.", "Lancer mops up; stands on A.")
phase("End of my turn — Vital Link.", "Control D + E (non-home) → 4 VP (counted). End-of-turn central: still "
  "his. I now hold B + D + E + A(his home) and he holds C → I control MORE objectives, denying his 6-VP swing next "
  "turn. Secondary: Behind Enemy Lines + Engage → 6 VP.")
tally("39", "31", "Primary 4 + Secondary 6. Crucially, I’ve flipped the objective COUNT — his Destroyer’s "
  "Wrath 6-VP clause is now denied, and I’m sitting on his home for the closing 10.")

turnhead("▶ His Turn 4 — Imperial Fists (desperate for kills)", FIST)
phase("Command.", "Gain 1 CP. Oath: the Lancer (it’s on his home — must be evicted). Discipline: Divination. "
  "He spends CP freely now: Armour of Contempt is defensive (useless on offence); he uses Fury of the First "
  "(stratagem) on the cyclone brick and Disciplined Extermination (+1 AP, Ignore Cover) to try to crack the Lancer.")
phase("Movement.", "The TH/SS brick finally turns for my Knights, but it’s M5 and a board away — it lumbers "
  "toward C’s Immolator. Cyclone Terminators reposition to shoot the Lancer on A.")
phase("Shooting.", "Everything into the Oathed Lancer on his home: cyclone brick (Fury +1 Hit, Oath, Fusillade, "
  "Disciplined Extermination +1 AP/Ignore Cover, Fury-of-the-First strat +1 Hit/+1 Wound) ≈ 14 raw → 4++ full "
  "+ FNP 6+ → ~9 → Lancer to ~11 W — still standing on his home.")
phase("Charge.", "Brick reaches the Immolator on C → charges it. Terminatus Assault battle-shocks the Sisters "
  "inside on disembark risk (they hold). Brick + Lysander → Immolator (T9 W11) explodes (Deadly Demise nothing "
  "— it’s an ally transport). Sisters melta half disembarks pre-charge and survive to contest.")
phase("Fight.", "Immolator dies; his brick is now stranded on C, far from my scoring, having spent two turns "
  "crossing the board to kill a transport.")
phase("End of his turn — Destroyer’s Wrath.", "Destroyed 1 unit (Immolator) → 3 VP. Objective clauses: I "
  "hold MORE objectives (B+D+E+A vs his C) → he gets NO 6-VP swing and no non-home 4 (he only has central C, but I "
  "contest… rule him C → 4 VP). ‘Destroyed more than I lost last turn’ (I lost 0 last turn) → 4 VP. "
  "Secondary: Bring It Down (Immolator) → 3 VP.")
tally("39", "45", "He’s ahead on the raw tally — Destroyer’s Wrath + Bring It Down farm my ablative pieces "
  "well. But he’s about to lose the two clauses that matter most: my home is safe, and I own HIS home.")

# ============================================================ ROUND 5
H("Battle Round 5 — the close", 1)
turnhead("▶ My Turn 5 — Knights (lock it down)", STEEL)
phase("Command.", "Gain 1 CP. Regen tops the Lancer (11→12) and Castellan. Vital Link round-2+ at end of "
  "Command: I control D + E + A (his home) — non-home → 4 VP. His brick squats C so no central bonus, but I "
  "have three-plus objectives to his one.")
phase("Movement.", "Consolidate every objective I can hold. Lancer STAYS on his home A. Crusader + Castellan + "
  "Helverin + Warglaive spread across B/D/E and box out his remaining mobile bodies. Sisters foot half never left B "
  "(home locked all game under the dome).")
phase("Shooting.", "Finish the cyclone Terminators (deny him a late scorer) and thin the Vanguard. I do not "
  "over-commit into the brick — it can’t reach anything that scores in one more turn.")
phase("Charge / Fight.", "Positional only — hold ranges, stay on objectives.")
phase("End of my turn — Vital Link.", "End-of-turn: control central? No (his brick). But this is the FIFTH "
  "round, so the round-2+ non-home clause also resolves at end of turn: D+E+A → 4 VP. Maintain Control marker on E "
  "with a spare unit → +1. Round primary ≈ 5. Secondary: Behind Enemy Lines (Lancer in his DZ) + Engage → "
  "6 VP.")
tally("50", "45", "Primary ~5 + Secondary 6. Now the END-OF-BATTLE clause is pending: I control his home.")

turnhead("▶ His Turn 5 — Imperial Fists (last swing)", FIST)
phase("Command.", "Gain 1 CP. Oath: the Lancer on his home (evict it or lose 10 VP). Discipline: Divination.")
phase("Movement / Shooting / Charge / Fight.", "He throws everything left at the Lancer on A — cyclone remnant "
  "shooting + the Bladeguard Ancient charging in. Through the Lancer’s 4++ FULL + FNP 6+, the combined damage "
  "≈ 12 → the Lancer (12 W after regen) is on the brink… and here variance decides a swing: on the "
  "average line the Lancer survives on ~1–3 W and HOLDS his home; on a bad run it dies and he denies the 10. "
  "This is the single rolled moment the whole game was built to reach — and I built the Lancer (4++ full + FNP + "
  "regen, the toughest single model I own) precisely to win it. Average line: it lives.")
phase("End of his turn — Destroyer’s Wrath.", "Destroyed maybe 1 unit (a Sister/Armiger) → 3 VP; "
  "objective clauses mostly denied (I still hold more) → ~0–4. Secondary: ~3.")
tally("50", "51", "Before the end-of-battle bonus, it is a DEAD HEAT — exactly the knife-edge a good practice game "
  "should be.")

# ============================================================ CONCLUSION
H("Result & why", 1)
P("End of Battle — Vital Link end-game clause:", bold=True, color=STEEL)
P("I control my opponent’s home objective A (the Lancer held it) → +10 VP.", size=11)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(6)
r = p.add_run("FINAL:  Knights 60  —  Imperial Fists 51"); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = STEEL
P("(Illustrative average-line tally after the 10-VP home bonus; secondaries are capped/estimated. The margin "
  "lives entirely in that closing 10 and in the objective-count denial of rounds 4–5.)", italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)

H("The five decisions that won it", 2)
bullet("Killed the Sternguard on Turn 1.", bold_lead="1. ")
P("   ~10 of his ~34-damage Oath alpha was unsaveable Sternguard Dev Wounds. Removing that leg on turn one meant "
  "his big turn-2 strike into the Castellan did ~14, not ~34 — it bracketed a Knight instead of killing one. "
  "Everything downstream flowed from that.", size=9)
bullet("Chose Freeblade Company for FNP.", bold_lead="2. ")
P("   Feel No Pain 6+ is the ONLY mechanic that touches his unsaveable Dev Wounds, and the 1-wound/turn regen "
  "un-did his chip. Rotate Ion Shields (4++) would not have — invulns don’t stop mortal-style wounds.", size=9)
bullet("Never fed the brick a premium Knight.", bold_lead="3. ")
P("   Screened with Armigers; the ~50-damage Oathed melee brick spent turns 2–4 killing a 140-pt bait Armiger "
  "and an ally transport, a board-corner away from anything that scored.", size=9)
bullet("Picked the objective mission, not the kill race.", bold_lead="4. ")
P("   Priority Assets / Vital Link let my speed + OC10 win the board while his two OC22 bricks could only squat one "
  "objective each. I conceded the centre and won B + D + E + his home instead — and denied his 6-VP "
  "‘more-objectives’ swing in rounds 4–5.", size=9)
bullet("Built the Lancer to win the last roll.", bold_lead="5. ")
P("   A 4++-full, FNP, regenerating M14 Knight is the toughest single model I own and the fastest — the perfect "
  "piece to sit on his home under fire for the closing 10 VP.", size=9)

H("Honest failure modes (where this flips)", 2)
bullet("If he had reserved the Sternguard instead of deploying them, I couldn’t alpha them T1 — they’d "
  "teleport in later at full strength and the ~34 alpha would be live. Counter: the Navigator dome + spreading Knights "
  "so a fresh volley can’t focus a bracketed one; kill them the turn they land.")
bullet("If I mis-screen and the brick reaches a big Knight, that’s ~50 Oathed damage — a dead premium Knight "
  "and a Bring-It-Down + Assassination bonus for him. One positioning error can cost the game.")
bullet("The Turn-5 Lancer-holds-the-home roll is genuine variance. If it dies, the 10-VP bonus vanishes and this is "
  "a loss. Bringing a second body to his home earlier (an Armiger) is the insurance if a real game allows it.")
bullet("His OC22 Terminators are better board-control than ‘everything’s OC1’ implied. If both bricks reach "
  "objectives, the objective count tightens — speed and killing his soft scorers (Intercessors, Vanguard) is what "
  "keeps me ahead on the count.")

P()
P("Bottom line: this is a ~55/45 matchup in the Knights’ favour — winnable, not free. It is decided by "
  "sequencing (kill the soft leg first), discipline (never feed the brick), and the mission (out-hold, steal the "
  "home) — exactly the parts of 40k that live above the dice.", bold=True, color=STEEL)

out = "docs/Great-Value-vs-Knights-Full-Game-Simulation.docx"
doc.save(out)
print("WROTE", out)
